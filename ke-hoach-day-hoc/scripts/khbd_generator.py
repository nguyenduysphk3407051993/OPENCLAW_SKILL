#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ke-hoach-day-hoc: Cong cu tao Ke Hoach Bai Day (KHBD) chuan CV5512
Tao file Word (.docx) day du cau truc cho mon KHTN lop 9 - bo KNTT.

PORTABILITY NOTE:
  Script nay KHONG co duong dan co dinh. Tat ca duong dan deu duoc
  truyen qua tham so dong lenh (CLI arguments). Agent can:
  1. Dung lenh 'locate' de tim duong dan tuyet doi cua script nay.
  2. Hoi nguoi dung thu muc KHBD mau va thu muc dau ra.
  3. Truyen cac duong dan do vao cac subcommand.

Subcommands:
  locate       - In ra duong dan tuyet doi cua script nay (de su dung portable)
  generate     - Tao file Word KHBD
  analyze      - Phan tich file KHBD mau, trich xuat noi dung ra JSON
  list-samples - Liet ke cac file KHBD mau trong thu muc

Usage:
  python khbd_generator.py locate
  python khbd_generator.py generate --ten-bai "Bai 21..." --mon-hoc "Hoa hoc" --so-tiet 5 --lop 9 --output out.docx
  python khbd_generator.py analyze --input sample.docx --output analysis.json
  python khbd_generator.py list-samples --ref-dir "/path/to/khbd/samples" --output list.json
"""

# /// script
# requires-python = ">=3.9"
# dependencies = [
#   "python-docx>=1.1.0",
# ]
# ///

import argparse
import json
import os
import sys
from pathlib import Path


def _setup_encoding():
    """Setup UTF-8 encoding for stdout/stderr on Windows."""
    import io
    if hasattr(sys.stdout, 'reconfigure'):
        try:
            sys.stdout.reconfigure(encoding='utf-8')
            sys.stderr.reconfigure(encoding='utf-8')
        except Exception:
            pass


def _get_docx():
    """Import python-docx components."""
    try:
        import docx
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        return docx, Document, Pt, Cm, RGBColor, Inches, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, qn, OxmlElement
    except ImportError as e:
        print(f"ERROR: Thiếu thư viện python-docx. Chạy: pip install python-docx\n{e}", file=sys.stderr)
        sys.exit(1)


def set_cell_background(cell, color_hex: str):
    """Set background color for a table cell."""
    _, _, _, _, _, _, _, _, _, qn, OxmlElement = _get_docx()
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_borders(table):
    """Apply borders to all cells in table."""
    _, _, _, _, _, _, _, _, _, qn, OxmlElement = _get_docx()
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_paragraph_with_style(doc, text: str, bold: bool = False, italic: bool = False,
                               font_size: int = 12, alignment=None, left_indent: float = 0):
    """Add a paragraph with formatting."""
    _, _, Pt, Cm, _, _, WD_ALIGN_PARAGRAPH, _, _, _, _ = _get_docx()
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    if alignment:
        para.alignment = alignment
    if left_indent:
        para.paragraph_format.left_indent = Cm(left_indent)
    return para


def create_activity_table(doc, rows_data: list[dict]):
    """
    Create a 2-column table (GV | HS) for an activity.
    rows_data: list of {"gv": "...", "hs": "..."}
    """
    _, _, Pt, _, _, _, _, _, WD_ALIGN_VERTICAL, _, _ = _get_docx()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False

    # Set column widths (roughly 50/50)
    for cell in table.columns[0].cells:
        cell.width = int(8.5 * 914400 / 25.4)  # ~8.5cm
    for cell in table.columns[1].cells:
        cell.width = int(8.5 * 914400 / 25.4)

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Hoạt động của GV'
    hdr_cells[1].text = 'Hoạt động của HS'

    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(11)
        set_cell_background(cell, 'DBEAFE')  # Light blue
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for row_data in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data.get('gv', '')
        row_cells[1].text = row_data.get('hs', '')
        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    set_cell_borders(table)
    return table


def build_default_content(ten_bai: str, mon_hoc: str, lop: int, so_tiet: int,
                            ten_gv: str = '', truong: str = '', to_bo_mon: str = '') -> dict:
    """Build default KHBD content when no JSON is provided."""
    # Determine subject-specific content
    if 'hóa' in mon_hoc.lower() or 'hoa' in mon_hoc.lower():
        subject_code = 'hoa'
        nang_luc_kien_thuc = [
            f"- Nêu được các kiến thức cơ bản về {ten_bai}.",
            "- Trình bày được các khái niệm, định nghĩa liên quan.",
            "- Giải thích được các hiện tượng hóa học liên quan đến bài học.",
            "- Viết được phương trình hóa học minh họa.",
        ]
        nang_luc_khtn = [
            "- Nhận thức khoa học tự nhiên: Nêu được các khái niệm, tính chất hóa học cơ bản.",
            "- Tìm hiểu tự nhiên: Quan sát thí nghiệm, rút ra kết luận về tính chất hóa học.",
            "- Vận dụng kiến thức: Giải thích hiện tượng thực tế, vận dụng vào bài tập.",
        ]
        thiet_bi_dung_cu = [
            "Máy chiếu, bảng nhóm, phiếu học tập.",
            "Ống nghiệm, giá đỡ ống nghiệm, kẹp gỗ, đèn cồn, ống nhỏ giọt.",
        ]
        thiet_bi_hoa_chat = [
            "Các hóa chất cần thiết theo nội dung thí nghiệm của bài.",
        ]
        phuong_phap = [
            "- Dạy học theo nhóm, nhóm cặp đôi.",
            "- Dạy học trực quan (thí nghiệm, hình ảnh, video).",
            "- Kĩ thuật công não, động não.",
            "- Dạy học nêu và giải quyết vấn đề thông qua câu hỏi trong SGK.",
        ]
    elif 'vật' in mon_hoc.lower() or 'vat' in mon_hoc.lower() or 'lí' in mon_hoc.lower():
        subject_code = 'vat_li'
        nang_luc_kien_thuc = [
            f"- Nêu được các kiến thức vật lí cơ bản về {ten_bai}.",
            "- Phát biểu được định nghĩa, công thức và đơn vị của các đại lượng liên quan.",
            "- Giải thích được các hiện tượng vật lí trong thực tế.",
            "- Vận dụng được công thức để giải bài tập định lượng.",
        ]
        nang_luc_khtn = [
            "- Nhận thức khoa học tự nhiên: Nêu được định nghĩa, công thức, đơn vị đo.",
            "- Tìm hiểu tự nhiên: Tiến hành thí nghiệm, thu thập số liệu, rút ra kết luận.",
            "- Vận dụng kiến thức: Giải bài tập, giải thích hiện tượng đời sống kỹ thuật.",
        ]
        thiet_bi_dung_cu = [
            "Máy chiếu, bảng nhóm, phiếu học tập.",
            "Bộ dụng cụ thí nghiệm theo nội dung bài học.",
        ]
        thiet_bi_hoa_chat = []
        phuong_phap = [
            "- Dạy học trực quan (thí nghiệm, hình ảnh minh họa).",
            "- Dạy học theo nhóm, hợp tác.",
            "- Dạy học nêu và giải quyết vấn đề.",
            "- Kĩ thuật khăn trải bàn, mảnh ghép.",
        ]
    else:  # Sinh học
        subject_code = 'sinh'
        nang_luc_kien_thuc = [
            f"- Nêu được các khái niệm cơ bản về {ten_bai}.",
            "- Trình bày được cơ chế, quá trình sinh học liên quan.",
            "- Phân tích được mối quan hệ giữa cấu trúc và chức năng.",
            "- Giải thích được các hiện tượng sinh học trong thực tế.",
        ]
        nang_luc_khtn = [
            "- Nhận thức khoa học tự nhiên: Nêu được khái niệm, đặc điểm cơ bản.",
            "- Tìm hiểu tự nhiên: Quan sát, phân tích hình vẽ, sơ đồ, rút ra kết luận.",
            "- Vận dụng kiến thức: Giải thích hiện tượng di truyền, liên hệ thực tiễn.",
        ]
        thiet_bi_dung_cu = [
            "Máy chiếu, bảng nhóm, phiếu học tập.",
            "Tranh ảnh, mô hình, sơ đồ liên quan đến bài học.",
        ]
        thiet_bi_hoa_chat = []
        phuong_phap = [
            "- Dạy học theo nhóm, nhóm cặp đôi.",
            "- Dạy học trực quan (hình vẽ, sơ đồ, mô hình).",
            "- Kĩ thuật công não, bản đồ tư duy.",
            "- Dạy học nêu và giải quyết vấn đề.",
        ]

    return {
        "truong": truong,
        "to_bo_mon": to_bo_mon,
        "ten_gv": ten_gv,
        "ten_bai": ten_bai,
        "mon_hoc": mon_hoc,
        "lop": lop,
        "so_tiet": so_tiet,
        "muc_tieu": {
            "kien_thuc": nang_luc_kien_thuc,
            "nang_luc_chung": [
                "- Tự chủ và tự học: Chủ động tìm kiếm thông tin, đọc SGK để hoàn thành nhiệm vụ học tập.",
                "- Giao tiếp và hợp tác: Hoạt động nhóm hiệu quả, đảm bảo các thành viên đều được tham gia.",
                "- Giải quyết vấn đề và sáng tạo: Giải quyết vấn đề kịp thời, thảo luận hiệu quả trong nhóm.",
            ],
            "nang_luc_khtn": nang_luc_khtn,
            "pham_chat": [
                "- Chăm chỉ: Chủ động tích cực đọc tài liệu, nghiên cứu SGK, hoàn thành bài tập đúng hạn.",
                "- Trách nhiệm: Chủ động hoàn thành các nhiệm vụ được giao khi làm việc nhóm.",
                "- Trung thực: Báo cáo trung thực kết quả quan sát, thí nghiệm và làm bài kiểm tra.",
            ],
        },
        "thiet_bi": {
            "dung_cu_hoa_chat": thiet_bi_dung_cu + thiet_bi_hoa_chat,
            "phieu_hoc_tap": [
                "PHT số 1: Phiếu khám phá kiến thức mới (hoàn thành trong hoạt động Hình thành kiến thức).",
                "PHT số 2: Phiếu luyện tập (hoàn thành trong hoạt động Luyện tập).",
            ],
        },
        "phuong_phap": phuong_phap,
        "hoat_dong": [
            {
                "so_thu_tu": 1,
                "ten": "Khởi động",
                "muc_tieu": (
                    f"Tạo hứng thú, kích thích tò mò của HS về nội dung {ten_bai}. "
                    "Kết nối kiến thức đã biết với bài học mới."
                ),
                "noi_dung": (
                    "GV đặt câu hỏi tình huống hoặc cho HS xem hình ảnh/video ngắn liên quan đến bài học. "
                    "HS chia sẻ hiểu biết ban đầu."
                ),
                "san_pham": (
                    "Câu trả lời của HS về hiểu biết ban đầu. "
                    "HS xác định được nhu cầu tìm hiểu về bài học mới."
                ),
                "to_chuc": [
                    {
                        "gv": "Giao nhiệm vụ:\nGV chiếu hình ảnh/đặt câu hỏi tình huống liên quan đến "
                              f"nội dung {ten_bai}. Yêu cầu HS suy nghĩ và chia sẻ hiểu biết ban đầu.",
                        "hs": "HS quan sát hình ảnh/nghe câu hỏi, suy nghĩ cá nhân về vấn đề GV đặt ra.",
                    },
                    {
                        "gv": "Hướng dẫn HS thực hiện nhiệm vụ:\nGV quan sát, khuyến khích HS phát biểu. "
                              "Không nhận xét đúng/sai ngay.",
                        "hs": "HS thảo luận cặp đôi (1-2 phút), chuẩn bị câu trả lời.",
                    },
                    {
                        "gv": "Báo cáo, thảo luận:\nGV mời 2-3 HS trình bày ý kiến. "
                              "Ghi nhận các câu trả lời lên bảng.",
                        "hs": "HS trình bày ý kiến cá nhân. Các HS khác lắng nghe, bổ sung.",
                    },
                    {
                        "gv": "Kết luận, nhận định:\nGV nhận xét chung, chưa kết luận đúng/sai. "
                              "Dẫn dắt: 'Để trả lời chính xác, chúng ta cùng tìm hiểu bài học hôm nay...'",
                        "hs": "HS lắng nghe, xác định vấn đề cần tìm hiểu trong bài học mới.",
                    },
                ],
            },
            {
                "so_thu_tu": 2,
                "ten": "Hình thành kiến thức mới",
                "muc_tieu": " ".join(nang_luc_kien_thuc[:2]),
                "noi_dung": (
                    "HS đọc SGK, nghiên cứu tài liệu, quan sát thí nghiệm/hình ảnh, "
                    "thảo luận nhóm và hoàn thành PHT số 1 để nắm vững nội dung chính của bài."
                ),
                "san_pham": (
                    "PHT số 1 được hoàn thành chính xác. "
                    "HS nắm được các kiến thức trọng tâm của bài học."
                ),
                "to_chuc": [
                    {
                        "gv": "Giao nhiệm vụ:\nGV chia lớp thành các nhóm 4-6 HS. "
                              "Phát PHT số 1. Yêu cầu các nhóm đọc SGK trang [...], "
                              "nghiên cứu nội dung và hoàn thành phiếu học tập trong 10 phút.",
                        "hs": "HS nhận PHT, phân công nhiệm vụ trong nhóm. "
                              "Đọc SGK và nghiên cứu nội dung được giao.",
                    },
                    {
                        "gv": "Hướng dẫn HS thực hiện nhiệm vụ:\nGV đi quan sát các nhóm, "
                              "hỗ trợ khi cần thiết. Đặt câu hỏi gợi mở nếu nhóm bị khó.",
                        "hs": "Các nhóm thảo luận, hoàn thành PHT. "
                              "Ghi kết quả vào phiếu hoặc bảng nhóm.",
                    },
                    {
                        "gv": "Báo cáo kết quả:\nGV mời đại diện 1-2 nhóm trình bày. "
                              "Các nhóm khác lắng nghe, nhận xét, bổ sung. "
                              "GV ghi nhận các ý kiến.",
                        "hs": "Đại diện nhóm trình bày kết quả. "
                              "Các nhóm khác nhận xét, bổ sung phần trình bày của nhóm bạn.",
                    },
                    {
                        "gv": "Tổng kết:\nGV chiếu/viết bảng nội dung kiến thức chuẩn. "
                              "Phân tích, làm rõ các điểm HS còn chưa nắm vững. "
                              "Chốt kiến thức trọng tâm.",
                        "hs": "HS lắng nghe, đối chiếu với phiếu học tập. "
                              "Ghi chép kiến thức vào vở.",
                    },
                ],
            },
            {
                "so_thu_tu": 3,
                "ten": "Luyện tập",
                "muc_tieu": (
                    "Củng cố và khắc sâu kiến thức vừa học thông qua việc giải bài tập. "
                    "Rèn luyện kỹ năng vận dụng kiến thức."
                ),
                "noi_dung": (
                    "HS làm bài tập trắc nghiệm và/hoặc tự luận ngắn về nội dung vừa học, "
                    "sử dụng PHT số 2 hoặc phần mềm trực tuyến (Quizziz/Kahoot)."
                ),
                "san_pham": (
                    "Đáp án bài tập chính xác. HS tự đánh giá được mức độ hiểu bài."
                ),
                "to_chuc": [
                    {
                        "gv": "Giao nhiệm vụ:\nGV phát PHT số 2 hoặc chiếu bài tập lên màn hình. "
                              "Yêu cầu HS làm bài cá nhân trong 5-7 phút.",
                        "hs": "HS nhận nhiệm vụ, làm bài tập cá nhân.",
                    },
                    {
                        "gv": "Hướng dẫn HS thực hiện nhiệm vụ:\nGV quan sát HS làm bài, "
                              "nhắc nhở HS tập trung. Không hỗ trợ trực tiếp để đánh giá thực chất.",
                        "hs": "HS làm bài tập độc lập. Ghi đáp án vào PHT.",
                    },
                    {
                        "gv": "Báo cáo kết quả:\nGV chiếu đáp án. "
                              "Mời 1-2 HS lên bảng trình bày cách giải (với bài tự luận). "
                              "Nhận xét, chữa bài.",
                        "hs": "HS đối chiếu bài làm với đáp án. "
                              "Sửa bài vào vở. Đặt câu hỏi nếu chưa hiểu.",
                    },
                    {
                        "gv": "Kết luận, nhận định:\nGV thống kê tỉ lệ HS làm đúng, "
                              "nhận xét ý thức làm bài. Nhấn mạnh lại các điểm dễ nhầm lẫn.",
                        "hs": "HS lắng nghe nhận xét. Ghi chú các lỗi thường gặp.",
                    },
                ],
            },
            {
                "so_thu_tu": 4,
                "ten": "Vận dụng",
                "muc_tieu": (
                    "Vận dụng kiến thức đã học vào giải quyết tình huống thực tế. "
                    "Phát triển năng lực tìm tòi, sáng tạo."
                ),
                "noi_dung": (
                    "HS giải quyết bài tập/câu hỏi vận dụng thực tế liên quan đến nội dung bài học. "
                    "Thực hiện ở nhà hoặc cuối buổi học."
                ),
                "san_pham": (
                    "Bài làm vận dụng nộp vào buổi học sau (hoặc cuối buổi học nếu còn thời gian). "
                    "Trình bày rõ ràng, có liên hệ với thực tiễn."
                ),
                "to_chuc": [
                    {
                        "gv": "Giao nhiệm vụ:\nGV nêu câu hỏi/bài tập vận dụng thực tế:\n"
                              "'Em hãy tìm hiểu và giải thích một hiện tượng/ứng dụng thực tế "
                              f"liên quan đến {ten_bai} trong cuộc sống hàng ngày.'\n"
                              "Yêu cầu HS hoàn thành và nộp vào buổi học sau.",
                        "hs": "HS ghi chép yêu cầu và câu hỏi/bài tập vào vở.",
                    },
                    {
                        "gv": "Hướng dẫn thực hiện:\nGV hướng dẫn HS cách tìm kiếm thông tin: "
                              "SGK, sách tham khảo, internet (các nguồn đáng tin cậy). "
                              "Gợi ý một số từ khóa tìm kiếm.",
                        "hs": "HS lắng nghe hướng dẫn. "
                              "Ghi chú các nguồn tài liệu tham khảo GV gợi ý.",
                    },
                    {
                        "gv": "Báo cáo kết quả (buổi học sau):\nGV mời 2-3 HS trình bày kết quả. "
                              "Nhận xét, đánh giá bài làm của HS.",
                        "hs": "HS trình bày bài làm vận dụng. "
                              "Các HS khác lắng nghe, đặt câu hỏi.",
                    },
                    {
                        "gv": "Kết luận, nhận định:\nGV nhận xét ý thức làm bài, "
                              "nhắc nhở HS không nộp bài hoặc nộp không đúng quy định (nếu có). "
                              "Dặn dò HS chuẩn bị bài học tiếp theo.",
                        "hs": "HS lắng nghe nhận xét. Ghi chú bài học tiếp theo cần chuẩn bị.",
                    },
                ],
            },
        ],
    }


def cmd_generate(args):
    """Generate KHBD Word document."""
    docx_module, Document, Pt, Cm, RGBColor, Inches, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, qn, OxmlElement = _get_docx()

    # Load content
    if args.content_json:
        with open(args.content_json, encoding='utf-8') as f:
            content = json.load(f)
        # Override with CLI args if provided
        if args.ten_bai:
            content['ten_bai'] = args.ten_bai
        if args.mon_hoc:
            content['mon_hoc'] = args.mon_hoc
        if args.so_tiet:
            content['so_tiet'] = args.so_tiet
        if args.lop:
            content['lop'] = args.lop
        if args.ten_gv:
            content['ten_gv'] = args.ten_gv
        if args.truong:
            content['truong'] = args.truong
        if args.to_bo_mon:
            content['to_bo_mon'] = args.to_bo_mon
    else:
        content = build_default_content(
            ten_bai=args.ten_bai,
            mon_hoc=args.mon_hoc,
            lop=args.lop,
            so_tiet=args.so_tiet,
            ten_gv=getattr(args, 'ten_gv', '') or '',
            truong=getattr(args, 'truong', '') or '',
            to_bo_mon=getattr(args, 'to_bo_mon', '') or '',
        )

    # Create document
    doc = Document()

    # Set page margins
    from docx.oxml.ns import qn as _qn
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # ==================== HEADER TABLE ====================
    header_table = doc.add_table(rows=2, cols=2)
    header_table.style = 'Table Grid'

    r0 = header_table.rows[0]
    r0.cells[0].text = f"Trường: {content.get('truong', '...........................')}"
    r0.cells[1].text = f"Họ và tên giáo viên: {content.get('ten_gv', '............................')}"

    r1 = header_table.rows[1]
    r1.cells[0].text = f"Tổ: {content.get('to_bo_mon', '................................')}"
    r1.cells[1].text = ""

    for row in header_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
    set_cell_borders(header_table)

    doc.add_paragraph()

    # ==================== TITLE ====================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("KHUNG KẾ HOẠCH BÀI DẠY")
    run.bold = True
    run.font.size = Pt(14)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("(Kèm theo Công văn số 5512/BGDĐT-GDTrH ngày 18 tháng 12 năm 2020 của Bộ GDĐT)")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Bài / Chủ đề
    bai_para = doc.add_paragraph()
    bai_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = bai_para.add_run(content['ten_bai'])
    run.bold = True
    run.font.size = Pt(13)

    tiet_para = doc.add_paragraph()
    tiet_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tiet_para.add_run(f"Môn học: {content['mon_hoc']} - Lớp {content['lop']}   |   "
                             f"Thời lượng: {content['so_tiet']} tiết")
    run.font.size = Pt(12)

    doc.add_paragraph()

    # ==================== I. MỤC TIÊU ====================
    h1 = doc.add_paragraph()
    run = h1.add_run("I. MỤC TIÊU")
    run.bold = True
    run.font.size = Pt(13)

    # 1. Kiến thức
    h2 = doc.add_paragraph()
    run = h2.add_run("1. Về kiến thức")
    run.bold = True
    run.font.size = Pt(12)

    for item in content['muc_tieu'].get('kien_thuc', []):
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)
        for run in p.runs:
            run.font.size = Pt(11)

    # 2. Năng lực
    h2 = doc.add_paragraph()
    run = h2.add_run("2. Về năng lực")
    run.bold = True
    run.font.size = Pt(12)

    h3 = doc.add_paragraph()
    run = h3.add_run("a) Năng lực chung")
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    h3.paragraph_format.left_indent = Cm(0.5)

    for item in content['muc_tieu'].get('nang_luc_chung', []):
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
        for run in p.runs:
            run.font.size = Pt(11)

    h3b = doc.add_paragraph()
    run = h3b.add_run("b) Năng lực khoa học tự nhiên")
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    h3b.paragraph_format.left_indent = Cm(0.5)

    for item in content['muc_tieu'].get('nang_luc_khtn', []):
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
        for run in p.runs:
            run.font.size = Pt(11)

    # 3. Phẩm chất
    h2 = doc.add_paragraph()
    run = h2.add_run("3. Về phẩm chất")
    run.bold = True
    run.font.size = Pt(12)

    for item in content['muc_tieu'].get('pham_chat', []):
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_paragraph()

    # ==================== II. THIẾT BỊ ====================
    h1 = doc.add_paragraph()
    run = h1.add_run("II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU")
    run.bold = True
    run.font.size = Pt(13)

    h2 = doc.add_paragraph()
    run = h2.add_run("Chuẩn bị dụng cụ, hóa chất")
    run.bold = True
    run.font.size = Pt(12)

    for item in content['thiet_bi'].get('dung_cu_hoa_chat', []):
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)
        for run in p.runs:
            run.font.size = Pt(11)

    h2 = doc.add_paragraph()
    run = h2.add_run("Phiếu học tập")
    run.bold = True
    run.font.size = Pt(12)

    for item in content['thiet_bi'].get('phieu_hoc_tap', []):
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_paragraph()

    # ==================== III. TIẾN TRÌNH ====================
    h1 = doc.add_paragraph()
    run = h1.add_run("III. TIẾN TRÌNH DẠY HỌC")
    run.bold = True
    run.font.size = Pt(13)

    # A. Phương pháp
    h2 = doc.add_paragraph()
    run = h2.add_run("A. PHƯƠNG PHÁP VÀ KĨ THUẬT DẠY HỌC")
    run.bold = True
    run.font.size = Pt(12)

    for item in content.get('phuong_phap', []):
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_paragraph()

    # B. Các hoạt động
    h2 = doc.add_paragraph()
    run = h2.add_run("B. CÁC HOẠT ĐỘNG DẠY HỌC")
    run.bold = True
    run.font.size = Pt(12)

    # Map activity number to label
    ten_hoat_dong_map = {
        1: "Khởi động",
        2: "Hình thành kiến thức mới",
        3: "Luyện tập",
        4: "Vận dụng",
    }

    for hd in content.get('hoat_dong', []):
        so_thu_tu = hd.get('so_thu_tu', 1)
        ten_hd = hd.get('ten', ten_hoat_dong_map.get(so_thu_tu, f"Hoạt động {so_thu_tu}"))

        doc.add_paragraph()

        # Activity header
        hd_header = doc.add_paragraph()
        run = hd_header.add_run(f"{so_thu_tu}. Hoạt động {so_thu_tu}: {ten_hd}")
        run.bold = True
        run.font.size = Pt(12)

        # a) Mục tiêu
        p = doc.add_paragraph()
        run = p.add_run("a) Mục tiêu: ")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(hd.get('muc_tieu', ''))
        run2.font.size = Pt(11)
        p.paragraph_format.left_indent = Cm(0.5)

        # b) Nội dung
        p = doc.add_paragraph()
        run = p.add_run("b) Nội dung: ")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(hd.get('noi_dung', ''))
        run2.font.size = Pt(11)
        p.paragraph_format.left_indent = Cm(0.5)

        # c) Sản phẩm
        p = doc.add_paragraph()
        run = p.add_run("c) Sản phẩm: ")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(hd.get('san_pham', ''))
        run2.font.size = Pt(11)
        p.paragraph_format.left_indent = Cm(0.5)

        # d) Tổ chức thực hiện
        p = doc.add_paragraph()
        run = p.add_run("d) Tổ chức thực hiện:")
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Cm(0.5)

        # Table GV | HS
        to_chuc = hd.get('to_chuc', [])
        if to_chuc:
            create_activity_table(doc, to_chuc)

    # ==================== SAVE ====================
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"[OK] Da tao file KHBD thanh cong!")
    print(f"[FILE] {output_path.resolve()}")


def cmd_analyze(args):
    """Analyze a sample KHBD Word file and extract content to JSON."""
    _, Document, Pt, _, _, _, _, _, _, _, _ = _get_docx()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: File không tồn tại: {input_path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(input_path))

    # Extract all text paragraphs
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraphs.append({
                "index": i,
                "style": para.style.name,
                "text": para.text.strip(),
            })

    # Extract table content
    tables = []
    for t_idx, table in enumerate(doc.tables):
        table_data = []
        for r_idx, row in enumerate(table.rows):
            row_data = []
            for c_idx, cell in enumerate(row.cells):
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        tables.append({"table_index": t_idx, "rows": table_data})

    result = {
        "source_file": str(input_path),
        "paragraphs": paragraphs,
        "tables": tables,
        "extracted_sections": _extract_sections(paragraphs, tables),
    }

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(str(output_path), 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"[OK] Phan tich hoan tat!")
    print(f"[INFO] Da trich xuat {len(paragraphs)} doan van va {len(tables)} bang.")
    print(f"[FILE] Ket qua: {output_path.resolve()}")


def _extract_sections(paragraphs: list, tables: list) -> dict:
    """Attempt to extract key sections from KHBD paragraphs."""
    sections = {
        "header": [],
        "muc_tieu": [],
        "thiet_bi": [],
        "tien_trinh": [],
        "hoat_dong": [],
    }
    current_section = "header"
    for para in paragraphs:
        text = para["text"].upper()
        if "MỤC TIÊU" in text or "I. MUC TIEU" in text:
            current_section = "muc_tieu"
        elif "THIẾT BỊ" in text or "HỌC LIỆU" in text or "II." in text:
            current_section = "thiet_bi"
        elif "TIẾN TRÌNH" in text or "III." in text:
            current_section = "tien_trinh"
        elif "HOẠT ĐỘNG" in text:
            current_section = "hoat_dong"

        sections[current_section].append(para["text"])

    return sections


def cmd_list_samples(args):
    """List all KHBD sample files in a directory."""
    ref_dir = Path(args.ref_dir)
    if not ref_dir.exists():
        print(f"ERROR: Thư mục không tồn tại: {ref_dir}", file=sys.stderr)
        sys.exit(1)

    docx_files = list(ref_dir.glob("*.docx"))
    docx_files = [f for f in docx_files if not f.name.startswith("~$")]  # Skip temp files

    result = {
        "directory": str(ref_dir),
        "total_files": len(docx_files),
        "files": [
            {
                "filename": f.name,
                "path": str(f),
                "size_kb": round(f.stat().st_size / 1024, 1),
            }
            for f in sorted(docx_files)
        ],
    }

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(str(output_path), 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"[OK] Da liet ke {len(docx_files)} file KHBD mau.")
    for f_info in result['files']:
        print(f"  - {f_info['filename']} ({f_info['size_kb']} KB)")
    print(f"[FILE] Danh sach day du: {output_path.resolve()}")


def cmd_locate():
    """
    In ra duong dan tuyet doi cua script nay.
    Agent dung lenh nay truoc khi goi cac subcommand khac
    de dam bao khong bi loi duong dan khi skill duoc chuyen sang moi truong moi.
    """
    script_path = Path(__file__).resolve()
    skill_dir = script_path.parent.parent  # ke_hoach_day_hoc/
    print(f"[SCRIPT] {script_path}")
    print(f"[SKILL_DIR] {skill_dir}")
    print(f"[REFERENCES] {skill_dir / 'references'}")
    print(f"[EXAMPLES] {skill_dir / 'examples'}")


def main():
    parser = argparse.ArgumentParser(
        description="Cong cu tao Ke Hoach Bai Day (KHBD) chuan CV5512 - file Word (.docx)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    subparsers = parser.add_subparsers(dest='command', required=True)

    # ---- locate ----
    subparsers.add_parser(
        'locate',
        help='In duong dan tuyet doi cua script (dung de goi portable tu bat ky thu muc nao)'
    )

    # ---- generate ----
    gen_parser = subparsers.add_parser('generate', help='Tao file Word KHBD')
    gen_parser.add_argument('--ten-bai', required=True, help='Ten bai hoc day du')
    gen_parser.add_argument('--mon-hoc', required=True,
                             choices=['Hoa hoc', 'Vat li', 'Sinh hoc',
                                      'Hóa học', 'Vật lí', 'Sinh học'],
                             help='Mon hoc: "Hoa hoc", "Vat li", hoac "Sinh hoc"')
    gen_parser.add_argument('--so-tiet', required=True, type=int, help='So tiet day (so nguyen)')
    gen_parser.add_argument('--lop', required=True, type=int, help='Lop hoc (vd: 9)')
    gen_parser.add_argument('--output', required=True, help='Duong dan file Word dau ra (.docx)')
    gen_parser.add_argument('--content-json', help='File JSON noi dung (tu lenh analyze)')
    gen_parser.add_argument('--ten-gv', default='', help='Ten giao vien')
    gen_parser.add_argument('--truong', default='', help='Ten truong')
    gen_parser.add_argument('--to-bo-mon', default='', help='Ten to bo mon')

    # ---- analyze ----
    ana_parser = subparsers.add_parser('analyze', help='Phan tich file KHBD mau (.docx) ra JSON')
    ana_parser.add_argument('--input', required=True, help='Duong dan file Word KHBD mau')
    ana_parser.add_argument('--output', required=True, help='Duong dan file JSON dau ra')

    # ---- list-samples ----
    ls_parser = subparsers.add_parser('list-samples', help='Liet ke file KHBD mau trong thu muc')
    ls_parser.add_argument('--ref-dir', required=True, help='Thu muc chua file KHBD mau')
    ls_parser.add_argument('--output', required=True, help='Duong dan file JSON danh sach')

    _setup_encoding()
    args = parser.parse_args()

    if args.command == 'locate':
        cmd_locate()
    elif args.command == 'generate':
        cmd_generate(args)
    elif args.command == 'analyze':
        cmd_analyze(args)
    elif args.command == 'list-samples':
        cmd_list_samples(args)


if __name__ == '__main__':
    main()
