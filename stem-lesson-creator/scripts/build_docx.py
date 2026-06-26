#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_docx.py — Render content.json thành Kế hoạch bài dạy STEM (.docx)

Cách dùng:
    python3 build_docx.py content.json "Giao_an_STEM.docx"

Yêu cầu: python-docx  (pip install python-docx --break-system-packages)
Định dạng: Times New Roman 12pt, heading đậm, bảng có nền tiêu đề xám,
phiếu học tập, rubric. Xem reference/schema.md để biết cấu trúc JSON.
"""
import sys, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
SIZE = 12
DOT = "." * 90  # dòng kẻ chấm trống

# ---------- low level helpers ----------
def _set_run(run, bold=False, italic=False, size=SIZE, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    # đảm bảo font áp dụng cho cả tiếng Việt (East Asian / complex script)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), FONT)
    if color is not None:
        run.font.color.rgb = color

def para(doc, text="", bold=False, italic=False, size=SIZE, align=None,
         left_indent=None, space_after=4, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)
    if text:
        r = p.add_run(text)
        _set_run(r, bold=bold, italic=italic, size=size, color=color)
    return p

def labeled(doc, label, text, left_indent=0.0, label_color=None):
    """Đoạn có nhãn đậm (vd 'GV:') + nội dung thường."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(4); pf.space_before = Pt(0); pf.line_spacing = 1.15
    if left_indent: pf.left_indent = Inches(left_indent)
    r1 = p.add_run(label + " ")
    _set_run(r1, bold=True, color=label_color)
    r2 = p.add_run(text)
    _set_run(r2)
    return p

def bullet(doc, text, mark="- ", indent=0.3):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(indent); pf.space_after = Pt(3)
    pf.space_before = Pt(0); pf.line_spacing = 1.15
    r = p.add_run(mark + text)
    _set_run(r)
    return p

def shade_cell(cell, fill="D9D9D9"):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcpr.append(shd)

def fill_cell(cell, text, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text if text is not None else "")
    _set_run(r, bold=bold)

def make_table(doc, headers, rows, widths=None):
    ncol = len(headers)
    t = doc.add_table(rows=1, cols=ncol)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        fill_cell(t.rows[0].cells[j], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(t.rows[0].cells[j])
    for row in rows:
        cells = t.add_row().cells
        for j in range(ncol):
            val = row[j] if j < len(row) else ""
            fill_cell(cells[j], val)
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    para(doc, "", space_after=2)
    return t

# ---------- section renderers ----------
def render_header(doc, meta):
    para(doc, meta.get("title", "KẾ HOẠCH BÀI DẠY STEM:"), bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=2)
    if meta.get("project_name"):
        para(doc, meta["project_name"], bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=6)
    for key, lbl in [("duration", "Thời lượng: "),
                     ("lead_subject", None), ("scale", None)]:
        pass
    if meta.get("duration"):
        para(doc, "Thời lượng: " + meta["duration"], bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    subj = ""
    if meta.get("lead_subject"):
        subj = "Môn học chủ trì: " + meta["lead_subject"]
        if meta.get("integrated_subjects"):
            subj += " | Môn học tích hợp: " + meta["integrated_subjects"] + "."
    if subj:
        para(doc, subj, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    if meta.get("scale"):
        para(doc, "Quy mô triển khai: " + meta["scale"], bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

def render_objectives(doc, obj):
    para(doc, "I. MỤC TIÊU BÀI HỌC", bold=True, size=13, space_after=4)
    kn = obj.get("knowledge", [])
    if kn:
        para(doc, "1. Kiến thức chuyên sâu", bold=True)
        for grp in kn:
            para(doc, grp.get("subject", ""), bold=True, left_indent=0.2)
            for it in grp.get("items", []):
                bullet(doc, it, indent=0.4)
    comp = obj.get("competencies", [])
    if comp:
        para(doc, "2. Định hướng phát triển năng lực", bold=True)
        for it in comp:
            bullet(doc, it)
    qual = obj.get("qualities", [])
    if qual:
        para(doc, "3. Phẩm chất", bold=True)
        for it in qual:
            bullet(doc, it)

def render_materials(doc, mat):
    para(doc, "II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU", bold=True, size=13, space_after=4)
    org = mat.get("organization", [])
    if org:
        para(doc, "1. Công tác tổ chức và chia nhóm (GV thực hiện trước tiết 1)",
             bold=True)
        for it in org:
            if it.strip().startswith("+"):
                bullet(doc, it.strip()[1:].strip(), mark="+ ", indent=0.55)
            elif it.strip().endswith(":"):
                para(doc, it, left_indent=0.3)
            else:
                bullet(doc, it)
    teacher = mat.get("teacher", [])
    groups = mat.get("groups", [])
    if teacher or groups:
        para(doc, "2. Danh mục học liệu và phân phối chuẩn bị", bold=True)
    if teacher:
        para(doc, "a. Giáo viên chuẩn bị:", bold=True, left_indent=0.2)
        for it in teacher:
            bullet(doc, it, indent=0.4)
    if groups:
        para(doc, "b. Mỗi nhóm học sinh tự chuẩn bị:", bold=True, left_indent=0.2)
        for it in groups:
            bullet(doc, it, indent=0.4)

def render_periods(doc, periods):
    para(doc, "III. TIẾN TRÌNH DẠY HỌC CHI TIẾT", bold=True, size=13, space_after=4)
    for per in periods:
        para(doc, per.get("title", ""), bold=True, size=12, space_after=4,
             color=RGBColor(0x1F, 0x3B, 0x73))
        if per.get("objectives"):
            para(doc, "1. Mục tiêu tiết học", bold=True)
            for it in per["objectives"]:
                bullet(doc, it)
        acts = per.get("activities", [])
        if acts:
            para(doc, "2. Tiến trình chi tiết", bold=True)
            for act in acts:
                para(doc, act.get("name", ""), bold=True, left_indent=0.15)
                for blk in act.get("blocks", []):
                    if "role" in blk:
                        labeled(doc, blk["role"] + ":", blk.get("text", ""),
                                left_indent=0.3,
                                label_color=RGBColor(0xC0, 0x00, 0x00)
                                if blk["role"] == "GV" else None)
                    elif blk.get("type") == "note":
                        labeled(doc, "Ghi chú:", blk.get("text", ""),
                                left_indent=0.3, label_color=RGBColor(0x55,0x55,0x55))
                    else:
                        para(doc, blk.get("text", ""), left_indent=0.3)

def render_formulas(doc, formulas):
    if not formulas: return
    para(doc, "IV. PHỤ LỤC CÔNG THỨC TOÁN HỌC", bold=True, size=13, space_after=4)
    for i, f in enumerate(formulas, 1):
        labeled(doc, f"{i}. {f.get('name','')}:", "", left_indent=0.0)
        if f.get("formula"):
            para(doc, f["formula"], italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=2)
        if f.get("explain"):
            para(doc, "Trong đó: " + f["explain"], left_indent=0.3, space_after=6)

def render_worksheets(doc, worksheets):
    if not worksheets: return
    para(doc, "V. HỆ THỐNG PHIẾU HỌC TẬP (NOTEBOOK) CHÍNH THỨC CỦA DỰ ÁN",
         bold=True, size=13, space_after=4)
    for ws in worksheets:
        para(doc, ws.get("title", ""), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             size=12, space_after=2)
        if ws.get("subtitle"):
            para(doc, ws["subtitle"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=4)
        for el in ws.get("elements", []):
            t = el.get("type")
            if t == "field":
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(el.get("label", "") + ": "); _set_run(r, bold=True)
                r2 = p.add_run("." * 60); _set_run(r2)
            elif t == "paragraph":
                para(doc, el.get("text", ""), space_after=4)
            elif t == "blanks":
                for _ in range(int(el.get("lines", 3))):
                    para(doc, DOT, space_after=2)
            elif t == "table":
                make_table(doc, el.get("headers", []), el.get("rows", []))
        para(doc, "", space_after=6)

def render_rubric(doc, rubric):
    if not rubric: return
    para(doc, "VI. TIÊU CHÍ ĐÁNH GIÁ TỔNG KẾT (RUBRIC ĐIỂM DỰ ÁN)",
         bold=True, size=13, space_after=4)
    make_table(doc, rubric.get("headers", []), rubric.get("rows", []))

# ---------- main ----------
def build(content, out_path):
    doc = Document()
    # margins & default style
    for s in doc.sections:
        s.left_margin = Inches(1.0); s.right_margin = Inches(0.79)
        s.top_margin = Inches(0.79); s.bottom_margin = Inches(0.79)
    normal = doc.styles['Normal']
    normal.font.name = FONT; normal.font.size = Pt(SIZE)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    render_header(doc, content.get("meta", {}))
    if content.get("objectives"): render_objectives(doc, content["objectives"])
    if content.get("materials"):  render_materials(doc, content["materials"])
    if content.get("periods"):    render_periods(doc, content["periods"])
    render_formulas(doc, content.get("formulas"))
    render_worksheets(doc, content.get("worksheets"))
    render_rubric(doc, content.get("rubric"))

    doc.save(out_path)
    return out_path

def main():
    if len(sys.argv) < 3:
        print("Cách dùng: python3 build_docx.py content.json output.docx")
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        content = json.load(f)
    out = build(content, sys.argv[2])
    print("Đã tạo:", out)

if __name__ == "__main__":
    main()
