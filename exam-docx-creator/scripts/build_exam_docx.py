#!/usr/bin/env python3
"""
build_exam_docx.py — Chuyển markdown câu hỏi giáo dục VN sang .docx đẹp mắt.

Replicate logic từ converter_service.py trong dự án WEB_APP_HO_TRO_SOAN_WORD:
- Page geometry: Letter, margins 2/1.25/1.5/1 cm
- Font: Times New Roman 12pt toàn bộ
- Question prefix: bold + blue (#0070C0)
- Options A-D: auto layout 4-col / 2-col / 1-per-line với tab stops
- Sub-items a-d: indent 0.5cm

Usage:
    python build_exam_docx.py --input questions.md --output de_thi.docx
    echo "Câu 1. ..." | python build_exam_docx.py --stdin --output de_thi.docx
    python build_exam_docx.py --input q.md --output q.docx --toggle-tex
    python build_exam_docx.py --input q.md --output q.docx --equation
"""
from __future__ import annotations

import argparse
import re
import sys
import unicodedata
from copy import deepcopy
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

# ============================================================
# Constants — giữ nguyên từ converter_service.py
# ============================================================

_QUESTION_PREFIX_RE = re.compile(
    r"^\s*((?:Câu|Bài|Ví dụ)\s+\d+[A-Za-z]?(?:\s*[:.)\-])?)",
    re.IGNORECASE,
)
_OPTION_PREFIX_RE = re.compile(r"^\s*([A-D])([.)])\s*(.+?)\s*$")
_SUBITEM_PREFIX_RE = re.compile(r"^\s*([a-d])([.)])\s*(.+?)\s*$")
_INLINE_STRUCTURE_RE = re.compile(
    r"(?<!\S)(?P<option>[A-D][.)])\s+|(?<!\S)(?P<subitem>[a-d][.)])\s+"
)

_DISPLAY_BRACKET_RE = re.compile(r"\\\[(.+?)\\\]", re.DOTALL)
_INLINE_PAREN_RE = re.compile(r"\\\((.+?)\\\)", re.DOTALL)
_DOLLAR_BLOCK_RE = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)
_DOLLAR_INLINE_RE = re.compile(
    r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)", re.DOTALL
)

_UNICODE_TEX_REPLACEMENTS = {
    "≤": r"\le ", "≥": r"\ge ", "≠": r"\ne ", "≈": r"\approx ",
    "×": r"\times ", "÷": r"\div ", "·": r"\cdot ", "±": r"\pm ",
    "∞": r"\infty ", "∈": r"\in ", "∉": r"\notin ",
    "⊂": r"\subset ", "⊆": r"\subseteq ", "∪": r"\cup ", "∩": r"\cap ",
    "∀": r"\forall ", "∃": r"\exists ",
    "→": r"\to ", "⇒": r"\Rightarrow ", "⇔": r"\Leftrightarrow ",
    "∑": r"\sum ", "∫": r"\int ", "√": r"\sqrt ",
    "π": r"\pi ", "α": r"\alpha ", "β": r"\beta ", "γ": r"\gamma ",
    "δ": r"\delta ", "Δ": r"\Delta ", "θ": r"\theta ",
    "λ": r"\lambda ", "μ": r"\mu ", "Ω": r"\Omega ",
    "–": "-", "—": "-", "−": "-",
}

_QUESTION_PREFIX_COLOR = RGBColor(0x00, 0x00, 0xFF)
_OPTION_PREFIX_COLOR = RGBColor(0x00, 0x00, 0xFF)
_SOLUTION_KEYWORDS_RE = re.compile(r"(Lời giải[:.]|Hướng dẫn giải[:.]|Đáp án[:.])", re.IGNORECASE)
_SOLUTION_KEYWORD_COLOR = RGBColor(0xC6, 0x0C, 0x4A)
_MD_SOLUTION_SPLIT_RE = re.compile(
    r"^([ \t*_]*((?:Lời giải|Hướng dẫn giải|Đáp án)[:.])[ \t*_]*)[ \t]*(?=\S)",
    re.IGNORECASE | re.MULTILINE
)

_DOCUMENT_FONT_NAME = "Times New Roman"
_DOCUMENT_BODY_FONT_SIZE = 12

_DOCUMENT_PAGE_WIDTH = Cm(21.59)
_DOCUMENT_PAGE_HEIGHT = Cm(27.94)
_DOCUMENT_LEFT_MARGIN = Cm(2.0)
_DOCUMENT_RIGHT_MARGIN = Cm(1.25)
_DOCUMENT_TOP_MARGIN = Cm(1.5)
_DOCUMENT_BOTTOM_MARGIN = Cm(1.0)

_OPTION_LEFT_INDENT = Cm(0.5)
_OPTION_TWO_COLUMN_TABS = (Cm(0.5), Cm(9.0))
_OPTION_FOUR_COLUMN_TABS = (Cm(0.5), Cm(4.5), Cm(9.0), Cm(13.0))
_OPTION_FOUR_COLUMN_MAX_LENGTH = 20
_OPTION_FOUR_COLUMN_TOTAL_LENGTH = 80
_OPTION_TWO_COLUMN_MAX_LENGTH = 40

_MARKDOWN_DOCX_FORMAT = (
    "markdown+raw_tex+tex_math_dollars+pipe_tables+grid_tables"
)
_MARKDOWN_TOGGLE_TEX_DOCX_FORMAT = (
    "markdown-raw_tex-tex_math_dollars-tex_math_single_backslash-"
    "tex_math_double_backslash+pipe_tables+grid_tables"
)


# ============================================================
# Sanitize TeX for Toggle TeX mode
# ============================================================

def sanitize_toggle_tex_latex(latex: str) -> str:
    """Loại Unicode/tiếng Việt trong vùng TeX để MathType Toggle TeX ổn định."""
    latex = latex.strip()
    for source, replacement in _UNICODE_TEX_REPLACEMENTS.items():
        latex = latex.replace(source, replacement)
    latex = latex.replace("đ", "d").replace("Đ", "D")
    latex = unicodedata.normalize("NFKD", latex)
    latex = "".join(char for char in latex if not unicodedata.combining(char))
    latex = latex.encode("ascii", "ignore").decode("ascii")
    latex = re.sub(r"[ \t]+", " ", latex)
    latex = re.sub(r"\s*\n\s*", "\n", latex)
    return latex.strip()


def normalize_toggle_tex_markdown(md: str) -> str:
    """Đưa delimiter về $...$/ \\[...\\] và làm sạch TeX."""
    # Cần gấp đôi dấu gạch chéo ngược (\\) thành (\\\\) để Pandoc không hiểu lầm là escape bracket mà in ra đúng một dấu gạch chéo ngược (\).
    md = _DOLLAR_BLOCK_RE.sub(
        lambda m: f"\\\\[\n{sanitize_toggle_tex_latex(m.group(1))}\n\\\\]", md
    )
    md = _DISPLAY_BRACKET_RE.sub(
        lambda m: f"\\\\[\n{sanitize_toggle_tex_latex(m.group(1))}\n\\\\]", md
    )
    md = _INLINE_PAREN_RE.sub(
        lambda m: f"${sanitize_toggle_tex_latex(m.group(1))}$", md
    )
    md = _DOLLAR_INLINE_RE.sub(
        lambda m: f"${sanitize_toggle_tex_latex(m.group(1))}$", md
    )
    return md


# ============================================================
# Markdown → DOCX (via pypandoc)
# ============================================================

def preprocess_markdown(md_content: str) -> str:
    """Tiền xử lý markdown: tách từ khóa lời giải và tự động thêm dòng trống trước danh sách."""
    # 1. Tách từ khóa lời giải/đáp án nếu viết chung dòng
    md_content = _MD_SOLUTION_SPLIT_RE.sub(r"\1\n\n", md_content)
    
    # 2. Tự động thêm dòng trống trước danh sách nếu bị thiếu (để Pandoc hiểu đúng)
    lines = md_content.splitlines()
    new_lines = []
    
    list_marker_re = re.compile(r"^\s*([-*•–—+]|a\)|b\)|c\)|d\)|\d+[.)])[ \t]+")
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped:
            is_list_item = bool(list_marker_re.match(line))
            if is_list_item and i > 0 and len(new_lines) > 0:
                prev_line = new_lines[-1].strip()
                if prev_line and not list_marker_re.match(new_lines[-1]):
                    new_lines.append("")
        new_lines.append(line)
        
    return "\n".join(new_lines)


def markdown_to_docx(md_content: str, output_path: Path, *, toggle_tex: bool = True) -> Path:
    """Chuyển markdown sang .docx rồi postprocess."""
    payload = md_content.strip() or " "
    payload = preprocess_markdown(payload)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if toggle_tex:
        payload = normalize_toggle_tex_markdown(payload)
        fmt = _MARKDOWN_TOGGLE_TEX_DOCX_FORMAT
    else:
        fmt = _MARKDOWN_DOCX_FORMAT

    pypandoc.convert_text(
        payload, "docx", format=fmt,
        outputfile=str(output_path),
        extra_args=["--wrap=preserve"],
    )
    postprocess_ocr_docx(output_path)
    return output_path


# ============================================================
# Postprocess — format câu hỏi, phương án, geometry, font
# ============================================================

def postprocess_ocr_docx(docx_path: Path) -> Path:
    """Hậu xử lý DOCX: định dạng câu hỏi, phương án, ý nhỏ."""
    document = Document(docx_path)
    _apply_document_geometry(document)
    _apply_document_font(document)
    original_paragraphs = list(document.paragraphs)

    for paragraph in original_paragraphs:
        if not paragraph.text or not paragraph.text.strip():
            continue

        segments = _explode_structured_paragraph(paragraph.text)
        current = paragraph

        if len(segments) > 1:
            _set_paragraph_text(current, segments[0][1])
            _apply_segment_format(current, segments[0][0])
            for kind, text in segments[1:]:
                current = _insert_paragraph_after(current, text)
                _apply_segment_format(current, kind)
            continue

        _set_paragraph_text(current, segments[0][1])
        _apply_segment_format(current, segments[0][0])

    _layout_answer_choices(document)
    _format_solution_sections(document)
    _format_tables(document)
    _format_images(document)
    
    for paragraph in document.paragraphs:
        _highlight_solution_keywords(paragraph)
        
    document.save(docx_path)
    return docx_path


# ============================================================
# Page geometry
# ============================================================

def _apply_document_geometry(document: Document) -> None:
    for section in document.sections:
        section.page_width = _DOCUMENT_PAGE_WIDTH
        section.page_height = _DOCUMENT_PAGE_HEIGHT
        section.left_margin = _DOCUMENT_LEFT_MARGIN
        section.right_margin = _DOCUMENT_RIGHT_MARGIN
        section.top_margin = _DOCUMENT_TOP_MARGIN
        section.bottom_margin = _DOCUMENT_BOTTOM_MARGIN


# ============================================================
# Font — Times New Roman toàn bộ
# ============================================================

def _apply_document_font(document: Document) -> None:
    for style in document.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        _set_font_family(style.font, style.element.get_or_add_rPr())

    normal_style = document.styles["Normal"]
    normal_style.font.size = Pt(_DOCUMENT_BODY_FONT_SIZE)

    for paragraph in _iter_document_paragraphs(document):
        for run in paragraph.runs:
            _set_run_font(run, size=None)


def _iter_document_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_document_paragraphs(cell)


def _set_font_family(font, r_pr) -> None:
    font.name = _DOCUMENT_FONT_NAME
    r_fonts = r_pr.get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{name}"), _DOCUMENT_FONT_NAME)


def _set_run_font(run, *, size=_DOCUMENT_BODY_FONT_SIZE, bold=None) -> None:
    r_pr = run._element.get_or_add_rPr()
    _set_font_family(run.font, r_pr)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


# ============================================================
# Paragraph structure: explode inline options/subitems
# ============================================================

def _explode_structured_paragraph(text: str) -> list[tuple[str, str]]:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text).strip()
    if not text:
        return [("paragraph", "")]

    segments: list[tuple[str, str]] = []
    question_match = _QUESTION_PREFIX_RE.match(text)

    if question_match:
        first_marker = _INLINE_STRUCTURE_RE.search(text, question_match.end())
        if not first_marker:
            return [("question", text)]
        question_text = text[:first_marker.start()].strip()
        if question_text:
            segments.append(("question", question_text))
        text = text[first_marker.start():].strip()

    markers = list(_INLINE_STRUCTURE_RE.finditer(text))
    if not markers:
        if segments:
            return segments
        if _OPTION_PREFIX_RE.match(text):
            return [("option", text)]
        if _SUBITEM_PREFIX_RE.match(text):
            return [("subitem", text)]
        return [("paragraph", text)]

    if not segments and markers[0].start() > 0:
        prefix_text = text[:markers[0].start()].strip()
        if prefix_text:
            segments.append(("paragraph", prefix_text))

    for index, marker in enumerate(markers):
        end = markers[index + 1].start() if index + 1 < len(markers) else len(text)
        chunk = text[marker.start():end].strip()
        if not chunk:
            continue
        kind = "option" if marker.group("option") else "subitem"
        segments.append((kind, chunk))

    return segments or [("paragraph", text)]


# ============================================================
# Segment formatting
# ============================================================

def _apply_segment_format(paragraph: Paragraph, kind: str) -> None:
    if kind == "question":
        _format_question_prefix(paragraph)
    elif kind == "option":
        _format_choice_paragraph(paragraph, is_subitem=False)
    elif kind == "subitem":
        _format_choice_paragraph(paragraph, is_subitem=True)


def _format_question_prefix(paragraph: Paragraph) -> None:
    text = paragraph.text.strip()
    match = _QUESTION_PREFIX_RE.match(text)
    if not match:
        return

    prefix = match.group(1).strip()
    remainder = text[match.end():].lstrip()
    _clear_paragraph(paragraph)

    prefix_run = paragraph.add_run(prefix)
    prefix_run.bold = True
    prefix_run.font.color.rgb = _QUESTION_PREFIX_COLOR

    if remainder:
        paragraph.add_run(f" {remainder}")


def _format_choice_paragraph(paragraph: Paragraph, *, is_subitem: bool) -> None:
    text = paragraph.text.strip()
    pattern = _SUBITEM_PREFIX_RE if is_subitem else _OPTION_PREFIX_RE
    match = pattern.match(text)
    if not match:
        return

    marker = "".join(match.groups()[:-1])
    content = match.group(len(match.groups())).strip()
    _clear_paragraph(paragraph)

    marker_run = paragraph.add_run(f"{marker} ")
    content_run = paragraph.add_run(content)

    fmt = paragraph.paragraph_format
    fmt.left_indent = _OPTION_LEFT_INDENT
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(3)
    fmt.space_after = Pt(3)
    fmt.line_spacing = 1
    _reset_tab_stops(paragraph)
    fmt.tab_stops.add_tab_stop(
        _OPTION_LEFT_INDENT, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES
    )
    _set_run_font(marker_run, bold=not is_subitem)
    if not is_subitem:
        marker_run.font.color.rgb = _OPTION_PREFIX_COLOR
    _set_run_font(content_run, bold=False)


# ============================================================
# Layout A-D answer choices (4-col / 2-col / 1-per-line)
# ============================================================

def _layout_answer_choices(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    index = 0

    while index < len(paragraphs):
        first = _parse_choice(paragraphs[index].text)
        if not first or first[0] != "A":
            index += 1
            continue

        group: list[tuple[Paragraph, str, str]] = []
        expected_markers = ("A", "B", "C", "D")
        for offset, expected in enumerate(expected_markers):
            if index + offset >= len(paragraphs):
                break
            parsed = _parse_choice(paragraphs[index + offset].text)
            if not parsed or parsed[0] != expected:
                break
            group.append((paragraphs[index + offset], parsed[0], parsed[1]))

        if len(group) != 4:
            index += 1
            continue

        choices = [(marker, content) for _, marker, content in group]
        first_element = group[0][0]._p

        if _use_four_choice_columns(choices):
            _set_choice_row(group[0][0], choices, _OPTION_FOUR_COLUMN_TABS)
            for paragraph, _, _ in group[1:]:
                _remove_paragraph(paragraph)
        elif _use_two_choice_columns(choices):
            _set_choice_row(group[0][0], choices[:2], _OPTION_TWO_COLUMN_TABS)
            _set_choice_row(group[2][0], choices[2:], _OPTION_TWO_COLUMN_TABS)
            _remove_paragraph(group[1][0])
            _remove_paragraph(group[3][0])
        else:
            index += 1
            continue

        paragraphs = list(document.paragraphs)
        index = next(
            (i for i, p in enumerate(paragraphs) if p._p is first_element), index
        ) + 1


def _parse_choice(text: str):
    match = _OPTION_PREFIX_RE.match((text or "").strip())
    if not match:
        return None
    return match.group(1), match.group(3).strip()


def _use_four_choice_columns(choices):
    lengths = [len(c) for _, c in choices]
    return max(lengths, default=0) <= _OPTION_FOUR_COLUMN_MAX_LENGTH and sum(lengths) <= _OPTION_FOUR_COLUMN_TOTAL_LENGTH


def _use_two_choice_columns(choices):
    lengths = [len(c) for _, c in choices]
    return max(lengths, default=0) <= _OPTION_TWO_COLUMN_MAX_LENGTH


def _set_choice_row(paragraph: Paragraph, choices, tab_positions) -> None:
    _clear_paragraph(paragraph)
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    fmt = paragraph.paragraph_format
    fmt.left_indent = _OPTION_LEFT_INDENT
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(3)
    fmt.space_after = Pt(3)
    fmt.line_spacing = 1
    _reset_tab_stops(paragraph)
    for position in tab_positions:
        fmt.tab_stops.add_tab_stop(position, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)

    for choice_index, (marker, content) in enumerate(choices):
        if choice_index:
            tab_run = paragraph.add_run("\t")
            _set_run_font(tab_run)
        marker_run = paragraph.add_run(f"{marker}. ")
        _set_run_font(marker_run, bold=True)
        marker_run.font.color.rgb = _OPTION_PREFIX_COLOR
        content_run = paragraph.add_run(content)
        _set_run_font(content_run, bold=False)


# ============================================================
# Paragraph helpers
# ============================================================

def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.text == text:
        return
    paragraph.text = text


def _clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag == qn("w:pPr"):
            continue
        paragraph._p.remove(child)


def _reset_tab_stops(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is not None:
        p_pr.remove(tabs)


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def _highlight_solution_keywords(paragraph: Paragraph) -> None:
    runs = list(paragraph.runs)
    for run in runs:
        text = run.text
        if not text:
            continue
        
        matches = list(_SOLUTION_KEYWORDS_RE.finditer(text))
        if not matches:
            continue
            
        p = paragraph._p
        run_element = run._element
        
        last_end = 0
        for match in matches:
            before_text = text[last_end:match.start()]
            if before_text:
                new_run = paragraph.add_run(before_text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                _set_run_font(new_run)
                run_element.addprevious(new_run._element)
                
            kw_text = match.group(1)
            kw_run = paragraph.add_run(kw_text)
            kw_run.bold = True
            kw_run.font.color.rgb = _SOLUTION_KEYWORD_COLOR
            _set_run_font(kw_run)
            run_element.addprevious(kw_run._element)
            
            last_end = match.end()
            
        after_text = text[last_end:]
        if after_text:
            new_run = paragraph.add_run(after_text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            _set_run_font(new_run)
            run_element.addprevious(new_run._element)
            
        p.remove(run_element)


def _format_solution_sections(document: Document) -> None:
    """Định dạng các đoạn văn trong phần Lời giải (enumerate, itemize)."""
    in_solution = False
    
    l1_re = re.compile(r"^([-*•–—]|a\)|b\)|c\)|d\)|\d+[.)])\s+(.*)$")
    l2_re = re.compile(r"^\s*([+])\s+(.*)$")
    l2_alt_re = re.compile(r"^\s{2,}([-*•–—]|a\)|b\)|c\)|d\)|\d+[.)])\s+(.*)$")
    conclusion_re = re.compile(r"^(=>|Chọn|Đáp án:)\s*(.*)$", re.IGNORECASE)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # Gặp tiền tố câu hỏi hoặc tiêu đề phần mới -> thoát khỏi block Lời giải
        if _QUESTION_PREFIX_RE.match(text) or text.startswith("#"):
            in_solution = False
            continue
            
        # Bắt đầu block Lời giải
        if _SOLUTION_KEYWORDS_RE.match(text):
            in_solution = True
            fmt = paragraph.paragraph_format
            fmt.space_before = Pt(6)
            fmt.space_after = Pt(3)
            fmt.line_spacing = 1.15
            fmt.keep_with_next = True
            continue
            
        if in_solution:
            fmt = paragraph.paragraph_format
            fmt.space_before = Pt(2)
            fmt.space_after = Pt(2)
            fmt.line_spacing = 1.15
            
            style_name = paragraph.style.name
            
            # 1. Định dạng danh sách mặc định của Word (nếu có numPr hoặc Style bắt đầu bằng List)
            p_pr = paragraph._p.get_or_add_pPr()
            num_pr = p_pr.find(qn("w:numPr"))
            is_native_list = (num_pr is not None) or style_name.startswith("List")
            
            if is_native_list:
                ilvl = 0
                if num_pr is not None:
                    ilvl_elem = num_pr.find(qn("w:ilvl"))
                    if ilvl_elem is not None:
                        try:
                            ilvl = int(ilvl_elem.get(qn("w:val"), "0"))
                        except ValueError:
                            pass
                else:
                    if "2" in style_name or "3" in style_name:
                        ilvl = 1
                
                if ilvl > 0:
                    fmt.left_indent = Cm(1.25)
                    fmt.first_line_indent = Cm(-0.35)
                else:
                    fmt.left_indent = Cm(0.75)
                    fmt.first_line_indent = Cm(-0.35)
                continue
                
            # 2. Định dạng danh sách tự viết thủ công
            l1_match = l1_re.match(text)
            l2_match = l2_re.match(text) or l2_alt_re.match(text)
            conclusion_match = conclusion_re.match(text)
            
            if l1_match:
                prefix = l1_match.group(1)
                fmt.left_indent = Cm(0.75)
                fmt.first_line_indent = Cm(-0.35)
                _reset_tab_stops(paragraph)
                fmt.tab_stops.add_tab_stop(Cm(0.75), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
                _convert_manual_prefix_to_tab(paragraph, prefix)
                
            elif l2_match:
                prefix = l2_match.group(1)
                fmt.left_indent = Cm(1.25)
                fmt.first_line_indent = Cm(-0.35)
                _reset_tab_stops(paragraph)
                fmt.tab_stops.add_tab_stop(Cm(1.25), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
                _convert_manual_prefix_to_tab(paragraph, prefix)
                
            elif conclusion_match:
                fmt.left_indent = Cm(0.75)
                fmt.first_line_indent = Pt(0)
                
            else:
                # Đoạn văn thường trong phần lời giải - căn thẳng hàng với nội dung Level 1
                fmt.left_indent = Cm(0.75)
                fmt.first_line_indent = Pt(0)


def _convert_manual_prefix_to_tab(paragraph: Paragraph, prefix: str) -> None:
    """Chuyển đổi khoảng trắng sau marker danh sách thủ công thành ký tự tab để căn chỉnh cột đẹp."""
    text_so_far = ""
    for run in paragraph.runs:
        if not run.text:
            continue
        text_so_far += run.text
        if text_so_far.lstrip().startswith(prefix):
            stripped = run.text.lstrip()
            if stripped.startswith(prefix + " "):
                lead_space_len = len(run.text) - len(stripped)
                lead_space = run.text[:lead_space_len]
                rest = stripped[len(prefix) + 1:]
                run.text = f"{lead_space}{prefix}\t{rest}"
                break
            elif stripped == prefix or (len(stripped) > len(prefix) and stripped.startswith(prefix) and stripped[len(prefix)] in (' ', '\t')):
                lead_space_len = len(run.text) - len(stripped)
                lead_space = run.text[:lead_space_len]
                rest = stripped[len(prefix):]
                if rest.startswith(" ") or rest.startswith("\t"):
                    rest = rest[1:]
                run.text = f"{lead_space}{prefix}\t{rest}"
                break
def _set_cell_background(cell, hex_color: str) -> None:
    """Tô màu nền cho một ô bảng (Cell)."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_margins(cell, *, top: int = 100, bottom: int = 100, left: int = 150, right: int = 150) -> None:
    """Thiết lập khoảng đệm trong của một ô bảng (cell padding) theo dxa."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _set_table_borders(table) -> None:
    """Thiết lập đường viền bảng phong cách hiện đại (chỉ có viền ngang, không có viền dọc)."""
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is not None:
        tblPr.remove(borders)
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))


def _format_tables(document: Document) -> None:
    """Định dạng tất cả các bảng biểu trong tài liệu cho chuyên nghiệp."""
    for table in document.tables:
        # Căn giữa bảng
        table.alignment = WD_TAB_ALIGNMENT.CENTER
        
        # Thiết lập viền bảng phong cách hiện đại
        try:
            _set_table_borders(table)
        except Exception as e:
            print(f"Warning: Lỗi vẽ viền bảng: {e}")
        
        # 1. Định dạng hàng đầu (Header)
        if len(table.rows) > 0:
            header_row = table.rows[0]
            trPr = header_row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            
            for cell in header_row.cells:
                # Màu nền xám nhẹ `#F2F2F2`
                _set_cell_background(cell, "F2F2F2")
                _set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Tiêu đề in đậm, cỡ chữ 11pt
                    for run in paragraph.runs:
                        run.bold = True
                        _set_run_font(run, size=11)
                        
        # 2. Định dạng các hàng nội dung
        for row in table.rows[1:]:
            trPr = row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            
            for cell in row.cells:
                _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_font(run, size=11)


def _format_images(document: Document) -> None:
    """Định dạng vị trí và kích thước hình ảnh để không bị tràn lề."""
    # 1. Căn giữa các đoạn văn chứa hình vẽ
    for paragraph in document.paragraphs:
        has_drawing = False
        for run in paragraph.runs:
            if run._element.find(qn('w:drawing')) is not None:
                has_drawing = True
                break
        if has_drawing:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            
    # 2. Giới hạn độ rộng tối đa của ảnh là 12cm để tránh tràn lề
    max_width = Cm(12.0)
    for shape in document.inline_shapes:
        try:
            if shape.width > max_width:
                ratio = max_width / shape.width
                shape.width = max_width
                shape.height = int(shape.height * ratio)
        except Exception:
            pass


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Chuyển markdown câu hỏi giáo dục VN → .docx đẹp mắt"
    )
    parser.add_argument("--input", "-i", help="File markdown đầu vào")
    parser.add_argument("--stdin", action="store_true", help="Đọc markdown từ stdin")
    parser.add_argument("--output", "-o", default="output.docx", help="File docx đầu ra")
    parser.add_argument(
        "--toggle-tex", action="store_true", default=True,
        help="Giữ công thức TeX text (MathType Toggle TeX) [mặc định]"
    )
    parser.add_argument(
        "--equation", action="store_true",
        help="Chuyển công thức sang OMML (Word Equation)"
    )
    args = parser.parse_args()

    if args.stdin:
        md_content = sys.stdin.read()
    elif args.input:
        md_content = Path(args.input).read_text(encoding="utf-8")
    else:
        parser.error("Phải chỉ định --input hoặc --stdin")
        return

    use_toggle_tex = not args.equation

    output_path = Path(args.output)
    markdown_to_docx(md_content, output_path, toggle_tex=use_toggle_tex)
    print(f"OK: {output_path.resolve()}")


if __name__ == "__main__":
    main()
