# -*- coding: utf-8 -*-
"""Tái tạo bang_ma_tran.png / bang_dac_ta_ma_tran.png thành file .docx
giữ nguyên cấu trúc gộp ô (merge) và các trường thông tin cố định.

Số mức độ đánh giá là tham số: 3 mức (Biết/Hiểu/Vận dụng) theo form gốc của
Bộ, hoặc 4 mức (thêm Liên hệ thực tế) như các bộ đề trong output/ đang dùng.
Bảng tự giãn số cột theo số mức độ."""

import argparse
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"


# ---------------------------------------------------------------- helpers
def set_borders(table):
    tbl = table._tbl
    pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    pr.append(borders)


def set_fixed_layout(table):
    pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    pr.append(layout)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def set_row_height(row, cm):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))
    h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)


def write(cell, text, bold=False, italic=False, align="center",
          size=10, sup=None, sup_size=7):
    """sup: phần ký tự mũ (chú thích chân trang) đặt sau text."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.name = FONT
        r.font.size = Pt(size)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if sup:
        r = p.add_run(sup)
        r.font.name = FONT
        r.font.size = Pt(sup_size)
        r.font.superscript = True
        r.bold = bold
        r.italic = italic
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def merge(table, r1, c1, r2, c2):
    return table.cell(r1, c1).merge(table.cell(r2, c2))


def new_doc(landscape_width_cm):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    sec.left_margin = sec.right_margin = Cm(1.5)
    sec.top_margin = sec.bottom_margin = Cm(1.2)
    return doc


def apply_widths(table, px, total_cm):
    total_px = sum(px)
    widths = [Cm(total_cm * p / total_px) for p in px]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
    return widths


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(12)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


BHV = ["Biết", "Hiểu", "Vận dụng"]
BHVL = ["Biết", "Hiểu", "Vận dụng", "Liên hệ"]
LEVEL_PRESETS = {3: BHV, 4: BHVL}
TOTAL_CM = 26.4

# Bề rộng tương đối (px) của một cột mức độ, tra theo nhãn.
LEVEL_PX = {"Biết": 34, "Hiểu": 38, "Vận dụng": 40, "Liên hệ": 38}

# Chân bảng khối "Tổng" của ma trận: điểm và tỷ lệ % theo từng mức độ.
TONG_PRESETS = {
    3: (("4,0", "3,0", "3,0"), ("40", "30", "30")),
    4: (("5,0", "3,0", "1,0", "1,0"), ("50", "30", "10", "10")),
}

# Bề rộng gốc của form 3 mức độ — giữ nguyên để chạy lại không làm lệch
# hai file mẫu đã có trong assets/.
PX_MA_TRAN_3 = [35, 71, 63,
                34, 38, 40, 34, 38, 39, 34, 38, 39, 34, 38, 40,
                34, 38, 39, 37]
PX_DAC_TA_3 = [30, 75, 65, 121,
               35, 40, 40, 35, 48, 41, 35, 41, 39, 34, 39, 41]


def level_px(fixed, levels, ngroup, tail=()):
    """px cho toàn bảng: cột cố định + ngroup nhóm mức độ + cột đuôi."""
    return list(fixed) + [LEVEL_PX[lv] for _ in range(ngroup)
                          for lv in levels] + list(tail)


def resolve_levels(levels):
    if levels is None:
        return BHV
    if isinstance(levels, int):
        return LEVEL_PRESETS[levels]
    return list(levels)


# ============================================================ 1. MA TRẬN
def build_ma_tran(path, levels=None):
    levels = resolve_levels(levels)
    L = len(levels)
    doc = new_doc(TOTAL_CM)
    title(doc, "1. MA TRẬN ĐỀ KIỂM TRA ĐỊNH KÌ")

    # 3 cột cố định + 4 nhóm loại câu hỏi + 1 nhóm "Tổng" + 1 cột tỷ lệ
    NROW, NCOL = 16, 3 + 5 * L + 1
    t = doc.add_table(rows=NROW, cols=NCOL)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t)
    set_fixed_layout(t)
    px = PX_MA_TRAN_3 if L == 3 else level_px([35, 71, 63], levels, 5, [37])
    apply_widths(t, px, TOTAL_CM)

    c_tong = 3 + 4 * L          # cột đầu khối "Tổng"
    c_tyle = 3 + 5 * L          # cột "Tỷ lệ % điểm"

    # ---- vùng tiêu đề (4 dòng)
    write(merge(t, 0, 0, 3, 0), "TT", bold=True)
    write(merge(t, 0, 1, 3, 1), "Chủ đề/\nChương", bold=True)
    write(merge(t, 0, 2, 3, 2), "Nội dung/đơn vị kiến thức", bold=True)
    write(merge(t, 0, 3, 0, c_tong - 1), "Mức độ đánh giá", bold=True)
    write(merge(t, 0, c_tong, 2, c_tyle - 1), "Tổng", bold=True)
    write(merge(t, 0, c_tyle, 3, c_tyle), "Tỷ lệ % điểm", bold=True)

    write(merge(t, 1, 3, 1, 3 + 3 * L - 1), "TNKQ", bold=True)
    write(merge(t, 1, 3 + 3 * L, 2, c_tong - 1), "Tự luận", bold=True)

    write(merge(t, 2, 3, 2, 3 + L - 1), "Nhiều lựa chọn", italic=True)
    write(merge(t, 2, 3 + L, 2, 3 + 2 * L - 1), "” Đúng – Sai” ",
          italic=True, sup="2")
    write(merge(t, 2, 3 + 2 * L, 2, 3 + 3 * L - 1), "Trả lời ngắn",
          italic=True, sup="3")

    for k, c in enumerate(range(3, c_tyle)):
        write(t.cell(3, c), levels[k % L], bold=True)

    for r in range(4):
        repeat_header(t.rows[r])
    set_row_height(t.rows[3], 1.0)

    # ---- phần thân: 3 chủ đề, mỗi chủ đề 3 dòng nội dung
    blocks = [("1", "Chủ đề 1"), ("2", "Chủ đề 2"), ("…", "Chủ đề …")]
    for i, (tt, cd) in enumerate(blocks):
        top = 4 + i * 3
        bot = top + 2
        write(merge(t, top, 0, bot, 0), tt)
        write(merge(t, top, 1, bot, 1), cd)
        for r in range(top, bot + 1):
            set_row_height(t.rows[r], 0.6)

    # ô ví dụ "(n)" ở cột "Đúng – Sai" / mức "Hiểu", dòng thứ 3 của Chủ đề 1
    write(t.cell(6, 3 + L + 1), "(n)", sup="4")

    # ---- phần chân bảng
    write(merge(t, 13, 0, 13, 2), "Tổng số câu/lệnh hỏi", bold=True, align="right")
    write(merge(t, 14, 0, 14, 2), "Tổng số điểm", bold=True, align="right")
    write(merge(t, 15, 0, 15, 2), "Tỷ lệ %", bold=True, align="right")

    tong_diem, tong_tyle = TONG_PRESETS[L]
    for g, (diem, tyle) in enumerate((("3,0", "30"), ("2,0", "20"),
                                      ("2,0", "20"), ("3,0", "30"))):
        c = 3 + g * L
        write(merge(t, 14, c, 14, c + L - 1), diem, sup="5" if g == 0 else None)
        write(merge(t, 15, c, 15, c + L - 1), tyle)

    for k in range(L):
        write(t.cell(14, c_tong + k), tong_diem[k])
        write(t.cell(15, c_tong + k), tong_tyle[k])

    for r in (13, 14, 15):
        set_row_height(t.rows[r], 0.55)

    doc.save(path)
    return NROW, NCOL


# ========================================================== 2. BẢNG ĐẶC TẢ
def build_dac_ta(path, levels=None):
    levels = resolve_levels(levels)
    L = len(levels)
    doc = new_doc(TOTAL_CM)
    title(doc, "2. BẢNG ĐẶC TẢ KIỂM TRA ĐỊNH KÌ")

    # thân: 3 khối chủ đề với số dòng lần lượt 10 / 9 / 9
    b1 = ["Biết:", "…", "", "", "Hiểu:", "…", "", "Vận dụng:", "…", ""]
    b2 = ["Biết:", "…", "", "Hiểu:", "…", "", "Vận dụng:", "…", ""]
    b3 = ["Biết:", "…", "", "", "Hiểu:", "…", "", "Vận dụng:", ""]
    blocks = [("1", "Chủ đề 1", b1), ("2", "Chủ đề 2", b2), ("…", "Chủ đề …", b3)]

    # 4 cột cố định + 4 nhóm loại câu hỏi (không có khối "Tổng")
    NCOL = 4 + 4 * L
    NROW = 4 + sum(len(b[2]) for b in blocks) + 3
    t = doc.add_table(rows=NROW, cols=NCOL)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t)
    set_fixed_layout(t)
    px = PX_DAC_TA_3 if L == 3 else level_px([30, 75, 65, 121], levels, 4)
    apply_widths(t, px, TOTAL_CM)

    # ---- vùng tiêu đề (4 dòng)
    write(merge(t, 0, 0, 3, 0), "TT", bold=True)
    write(merge(t, 0, 1, 3, 1), "Chủ đề/Chương", bold=True)
    write(merge(t, 0, 2, 3, 2), "Nội dung/đơn vị kiến thức", bold=True)
    write(merge(t, 0, 3, 3, 3), "Yêu cầu cần đạt", bold=True)
    write(merge(t, 0, 4, 0, NCOL - 1), "Số câu hỏi ở các mức độ đánh giá",
          bold=True)

    write(merge(t, 1, 4, 1, 4 + 3 * L - 1), "TNKQ", bold=True)
    write(merge(t, 1, 4 + 3 * L, 2, NCOL - 1), "Tự luận", bold=True)

    write(merge(t, 2, 4, 2, 4 + L - 1), "Nhiều lựa chọn", italic=True)
    write(merge(t, 2, 4 + L, 2, 4 + 2 * L - 1), "”Đúng – Sai”", italic=True)
    write(merge(t, 2, 4 + 2 * L, 2, 4 + 3 * L - 1), "Trả lời ngắn", italic=True)

    for k, c in enumerate(range(4, NCOL)):
        write(t.cell(3, c), levels[k % L], bold=True)

    for r in range(4):
        repeat_header(t.rows[r])
    set_row_height(t.rows[3], 1.0)

    # ---- phần thân
    r = 4
    for tt, cd, rows in blocks:
        top, bot = r, r + len(rows) - 1
        write(merge(t, top, 0, bot, 0), tt)
        write(merge(t, top, 1, bot, 1), cd)
        merge(t, top, 2, bot, 2)          # Nội dung/đơn vị kiến thức: gộp, để trống
        t.cell(top, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for k, txt in enumerate(rows):
            write(t.cell(top + k, 3), txt, bold=txt.endswith(":"), align="left")
            set_row_height(t.rows[top + k], 0.42)
        r = bot + 1

    # ô ví dụ "(n) (NL?)" ở cột "Đúng – Sai" / mức "Hiểu", dòng "Vận dụng:" của Chủ đề 2
    c_ds_hieu = 4 + L + 1
    r_vd_cd2 = 4 + len(b1) + b2.index("Vận dụng:")
    write(t.cell(r_vd_cd2, c_ds_hieu), "(n)\n(NL?)", sup="6")
    set_row_height(t.rows[r_vd_cd2], 0.72)

    # ---- phần chân bảng
    f0 = NROW - 3
    for k, lab in enumerate(("Tổng số câu", "Tổng số điểm", "Tỷ lệ %")):
        write(merge(t, f0 + k, 0, f0 + k, 2), lab, bold=True, align="right")
        set_row_height(t.rows[f0 + k], 0.42)

    for vals, rr in ((("3,0", "2,0", "2,0", "3,0"), f0 + 1),
                     (("30", "20", "20", "30"), f0 + 2)):
        for g, v in enumerate(vals):
            c = 4 + g * L
            write(merge(t, rr, c, rr, c + L - 1), v)

    doc.save(path)
    return NROW, NCOL


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--muc-do", type=int, choices=(3, 4), default=3,
                    help="số mức độ đánh giá (mặc định 3 = form gốc của Bộ)")
    ap.add_argument("-o", "--out-dir", default=os.path.dirname(os.path.abspath(__file__)),
                    help="thư mục xuất (mặc định: chính thư mục assets/)")
    ap.add_argument("--hau-to", default=None,
                    help="hậu tố tên file (mặc định: '' với 3 mức, '_4mucdo' với 4 mức)")
    a = ap.parse_args()

    levels = LEVEL_PRESETS[a.muc_do]
    hau_to = a.hau_to if a.hau_to is not None else ("" if a.muc_do == 3 else "_4mucdo")
    os.makedirs(a.out_dir, exist_ok=True)

    for ten, build in (("mau_bang_ma_tran", build_ma_tran),
                       ("mau_bang_dac_ta_ma_tran", build_dac_ta)):
        path = os.path.join(a.out_dir, f"{ten}{hau_to}.docx")
        nrow, ncol = build(path, levels)
        print(f"{os.path.basename(path):<38} {nrow} dòng x {ncol} cột"
              f"  ({' / '.join(levels)})")


if __name__ == "__main__":
    main()
