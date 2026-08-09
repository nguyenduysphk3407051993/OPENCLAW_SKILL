# -*- coding: utf-8 -*-
"""Tiện ích dựng file .docx giữ nguyên form mẫu (bang_ma_tran.docx /
bang_dac_ta_ma_tran.docx) cho skill exam-idea.

Mọi script khác trong skill đều import từ đây để bảo đảm 2 bảng luôn
đồng nhất về font, đường viền, độ rộng cột và cách gộp ô.
"""

from __future__ import annotations

import json
import os
import sys


def force_utf8_stdout():
    """Bảo đảm in được tiếng Việt trên VPS Linux chạy locale C/POSIX.

    Không có bước này, `print` chuỗi tiếng Việt sẽ ném UnicodeEncodeError khi
    biến môi trường LANG/LC_ALL chưa đặt UTF-8 (rất hay gặp trên VPS mới).
    """
    for stream in (sys.stdout, sys.stderr):
        enc = (getattr(stream, "encoding", None) or "").lower()
        if enc.replace("-", "") not in ("utf8", "utf_8") and hasattr(stream, "reconfigure"):
            try:
                stream.reconfigure(encoding="utf-8", errors="replace")
            except (ValueError, OSError):
                pass


force_utf8_stdout()

from docx import Document  # noqa: E402
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT = "Times New Roman"

# Bề ngang phần bảng (cm) trên khổ A4 ngang, lề 1.2cm
PAGE_W_CM = 29.7
PAGE_H_CM = 21.0
MARGIN_CM = 1.2
TABLE_W_CM = PAGE_W_CM - 2 * MARGIN_CM - 0.3


# --------------------------------------------------------------- low level
def _rfonts(run):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.append(rf)
    rf.set(qn("w:ascii"), FONT)
    rf.set(qn("w:hAnsi"), FONT)
    rf.set(qn("w:eastAsia"), FONT)
    rf.set(qn("w:cs"), FONT)


def set_borders(table, sz="6"):
    pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    pr.append(borders)


def set_fixed_layout(table):
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)


def set_cell_margins(table, left=40, right=40, top=10, bottom=10):
    """Lề trong ô, đơn vị dxa (1/20 pt)."""
    pr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("left", left),
                     ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    pr.append(mar)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def set_row_height(row, cm, rule="atLeast"):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))
    h.set(qn("w:hRule"), rule)
    trPr.append(h)


def shade(cell, hexcolor="F2F2F2"):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    tcPr.append(el)


ALIGN = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def write(cell, text="", bold=False, italic=False, align="center",
          size=9, sup=None, sup_size=6, valign="center"):
    """Ghi nội dung vào ô. `sup` là ký tự mũ (chú thích) nối sau `text`.

    Xuống dòng trong `text` dùng ký tự \n.
    """
    cell.vertical_alignment = {
        "center": WD_ALIGN_VERTICAL.CENTER,
        "top": WD_ALIGN_VERTICAL.TOP,
        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
    }[valign]
    p = cell.paragraphs[0]
    p.alignment = ALIGN[align]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if text not in (None, ""):
        for i, line in enumerate(str(text).split("\n")):
            if i:
                p.add_run().add_break()
            r = p.add_run(line)
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(size)
            _rfonts(r)
    if sup:
        r = p.add_run(str(sup))
        r.font.size = Pt(sup_size)
        r.font.superscript = True
        r.bold = bold
        r.italic = italic
        _rfonts(r)
    return cell


def merge(table, r1, c1, r2, c2):
    return table.cell(r1, c1).merge(table.cell(r2, c2))


def apply_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def new_doc(landscape=True, margin=MARGIN_CM, base_size=10):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(base_size)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(0)
    sec = doc.sections[0]
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Cm(PAGE_W_CM), Cm(PAGE_H_CM)
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width, sec.page_height = Cm(PAGE_H_CM), Cm(PAGE_W_CM)
    sec.left_margin = sec.right_margin = Cm(margin)
    sec.top_margin = sec.bottom_margin = Cm(margin)
    return doc


def para(doc, text="", bold=False, italic=False, align="left", size=11,
         space_before=0, space_after=2, indent=0.0):
    p = doc.add_paragraph()
    p.alignment = ALIGN[align]
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.left_indent = Cm(indent)
    for i, line in enumerate(str(text).split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        _rfonts(r)
    return p


def rich_para(doc, chunks, align="left", size=11, space_before=0,
              space_after=2, indent=0.0):
    """chunks: list các tuple (text, {bold, italic, size, sup})."""
    p = doc.add_paragraph()
    p.alignment = ALIGN[align]
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.left_indent = Cm(indent)
    for text, opt in chunks:
        r = p.add_run(str(text))
        r.bold = opt.get("bold", False)
        r.italic = opt.get("italic", False)
        r.font.size = Pt(opt.get("size", size))
        if opt.get("sup"):
            r.font.superscript = True
        _rfonts(r)
    return p


def title(doc, text, size=13, space_after=6):
    return para(doc, text, bold=True, align="center", size=size,
                space_after=space_after)


# ------------------------------------------------------------------ config
def save_doc(doc, path):
    """Lưu docx, tự tạo thư mục cha nếu chưa có (Linux không tự tạo)."""
    path = os.path.abspath(path)
    parent = os.path.dirname(path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    doc.save(path)
    return path


def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def dump_json(obj, path):
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)


def fmt_diem(x):
    """1.0 -> '1,0' ; 0.25 -> '0,25' ; 0 -> ''."""
    if x in (None, 0, 0.0):
        return ""
    s = f"{float(x):.2f}".rstrip("0").rstrip(".")
    if "." not in s:
        s += ".0"
    return s.replace(".", ",")


def fmt_int(x):
    return "" if not x else str(int(x))


def fmt_pct(x):
    if x in (None, 0, 0.0):
        return ""
    s = f"{float(x):.1f}".rstrip("0").rstrip(".")
    return s.replace(".", ",")


# --------------------------------------------------- cấu hình nhóm/mức độ
def get_muc_do(cfg):
    """Trả về list dict {ma, ten, ten_ngan}."""
    out = []
    for m in cfg["cau_hinh"]["muc_do"]:
        if isinstance(m, str):
            out.append({"ma": m, "ten": m, "ten_ngan": m})
        else:
            out.append({"ma": m.get("ma", m["ten"]),
                        "ten": m["ten"],
                        "ten_ngan": m.get("ten_ngan", m["ten"])})
    return out


def get_nhom(cfg):
    """Trả về list dict nhóm câu hỏi + gom theo nhóm cha (TNKQ / Tự luận)."""
    nhom = cfg["cau_hinh"]["nhom_cau_hoi"]
    parents = []
    for g in nhom:
        pc = g.get("nhom_cha", g["ten"])
        if not parents or parents[-1]["ten"] != pc:
            parents.append({"ten": pc, "con": []})
        parents[-1]["con"].append(g)
    return nhom, parents
