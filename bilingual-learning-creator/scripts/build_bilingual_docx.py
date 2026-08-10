#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Dựng tài liệu học tập song ngữ Việt–Anh (.docx) từ file JSON đúng schema.

Cách dùng:
    python scripts/build_bilingual_docx.py <input.json> -o <output.docx>

Xem cấu trúc JSON đầu vào ở assets/schema.json.
"""

import argparse
import json
import sys

# Đảm bảo in được tiếng Việt có dấu ra console (Windows mặc định cp1252).
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass

# ---------------------------------------------------------------------------
# Import python-docx với thông báo lỗi tiếng Việt rõ ràng
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "LỖI: Chưa cài thư viện 'python-docx'.\n"
        "Hãy cài bằng lệnh:\n\n"
        "    pip install python-docx\n\n"
    )
    sys.exit(2)


# ---------------------------------------------------------------------------
# Hằng số định dạng
# ---------------------------------------------------------------------------
FONT_NAME = "Times New Roman"
MONO_FONT = "Consolas"  # font monospace cho pattern ngữ pháp
FONT_SIZE = 13          # cỡ chữ mặc định (pt)
GRAY_FILL = "D9D9D9"    # nền xám nhạt cho header bảng / khung
LIGHT_GRAY = "F2F2F2"   # nền xám rất nhạt
GRAY_TEXT = (0x80, 0x80, 0x80)  # màu chữ xám cho ghi chú phụ

# Nhãn tiếng Việt cho từng loại cụm từ trong phrasebook
PHRASE_KIND_LABELS = {
    "collocation": "Cụm từ đi với nhau",
    "idiom": "Thành ngữ",
    "phrasal_verb": "Cụm động từ",
    "fixed_expression": "Cấu trúc cố định",
    "academic": "Ngôn ngữ học thuật",
    "classroom": "Ngôn ngữ lớp học",
}

# Nhãn tiếng Việt cho từ loại khi nhóm glossary theo pos
POS_LABELS = {
    "n": "n = danh từ",
    "v": "v = động từ",
    "adj": "adj = tính từ",
    "adv": "adv = trạng từ",
}


# ---------------------------------------------------------------------------
# Ngoại lệ validate
# ---------------------------------------------------------------------------
class ValidationError(Exception):
    """Lỗi dữ liệu JSON đầu vào không hợp lệ."""


# ---------------------------------------------------------------------------
# Tiện ích font: set eastAsia để tiếng Việt có dấu không vỡ font
# ---------------------------------------------------------------------------
def _apply_font(run, size=None, bold=None, italic=None, color=None):
    """Gán font Times New Roman cho run, đảm bảo cả w:eastAsia."""
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)

    run.font.size = Pt(size if size is not None else FONT_SIZE)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_run(paragraph, text, size=None, bold=None, italic=None, color=None):
    run = paragraph.add_run(text)
    _apply_font(run, size=size, bold=bold, italic=italic, color=color)
    return run


def add_mono_run(paragraph, text, size=None, bold=None, italic=None, color=None):
    """Thêm run dùng font monospace (Consolas) — cho pattern ngữ pháp."""
    run = paragraph.add_run(text)
    run.font.name = MONO_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), MONO_FONT)
    rfonts.set(qn("w:hAnsi"), MONO_FONT)
    rfonts.set(qn("w:cs"), MONO_FONT)
    rfonts.set(qn("w:eastAsia"), MONO_FONT)
    run.font.size = Pt(size if size is not None else FONT_SIZE)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    return run


def _set_cell_shading(cell, fill):
    """Tô nền cho ô bảng."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_para_shading(paragraph, fill):
    """Tô nền cho cả đoạn văn."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _remove_table_borders(table):
    """Bỏ toàn bộ viền của bảng."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def _fill_cell(cell, runs_spec, align=None):
    """Ghi nội dung vào ô bảng. runs_spec: list các dict kwargs cho add_run.

    Đoạn văn đầu tiên dùng luôn paragraph mặc định của ô.
    """
    cell.text = ""
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    for spec in runs_spec:
        text = spec.get("text", "")
        add_run(
            para,
            text,
            size=spec.get("size"),
            bold=spec.get("bold"),
            italic=spec.get("italic"),
            color=spec.get("color"),
        )
    return para


def add_thin_rule(doc):
    """Vẽ một đường kẻ ngang mảnh (viền dưới của một đoạn trống)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")       # mảnh
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")  # xám nhạt
    borders.append(bottom)
    p_pr.append(borders)
    return p


# ---------------------------------------------------------------------------
# Render dùng chung cho $defs.vocab và $defs.phrase
# ---------------------------------------------------------------------------
VOCAB_COLS = [
    ("word", "Từ"),
    ("pos", "Loại"),
    ("ipa", "IPA"),
    ("vi", "Nghĩa"),
    ("note_vi", "Ghi chú"),
]


def build_vocab_table(doc, vocab_items, size=11):
    """Dựng bảng nhỏ gọn cho danh sách $defs.vocab. Bỏ cột rỗng toàn bộ.

    Cột: Từ | Loại | IPA | Nghĩa | Ghi chú.
    """
    active = []
    for key, label in VOCAB_COLS:
        if key in ("word", "vi"):
            active.append((key, label))
        elif any(str(it.get(key, "")).strip() for it in vocab_items):
            active.append((key, label))

    table = doc.add_table(rows=1, cols=len(active))
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for c, (key, label) in enumerate(active):
        _fill_cell(hdr[c], [{"text": label, "bold": True, "size": size}],
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(hdr[c], GRAY_FILL)

    for it in vocab_items:
        row = table.add_row().cells
        for c, (key, label) in enumerate(active):
            val = str(it.get(key, "")).strip()
            italic = key == "ipa"
            _fill_cell(row[c], [{"text": val, "size": size, "italic": italic}])
    return table


def build_phrase_bullets(doc, phrases, indent=Cm(0.6), size=12):
    """Danh sách bullet cho $defs.phrase: 'en' đậm — 'vi', kèm ví dụ nếu có."""
    for ph in phrases:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = indent
        add_run(p, "• ", bold=True, size=size)
        add_run(p, str(ph.get("en", "")), bold=True, size=size)
        if ph.get("ipa"):
            add_run(p, f" {ph['ipa']}", italic=True, size=size - 1, color=GRAY_TEXT)
        add_run(p, f" — {ph.get('vi', '')}", size=size)
        ex_en = str(ph.get("example_en", "")).strip()
        ex_vi = str(ph.get("example_vi", "")).strip()
        if ex_en or ex_vi:
            pe = doc.add_paragraph()
            pe.paragraph_format.left_indent = indent + Cm(0.5)
            if ex_en:
                add_run(pe, f"Ví dụ: {ex_en}", italic=True, size=size - 1)
            if ex_vi:
                if ex_en:
                    add_run(pe, "  ", size=size - 1)
                add_run(pe, f"({ex_vi})", italic=True, size=size - 1, color=GRAY_TEXT)


# ---------------------------------------------------------------------------
# Thiết lập trang: A4, lề, font mặc định của style Normal
# ---------------------------------------------------------------------------
def setup_document(doc):
    # Font mặc định của style Normal
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)

    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)


# ---------------------------------------------------------------------------
# VALIDATE
# ---------------------------------------------------------------------------
VALID_TYPES = {
    "glossary", "parallel_text", "translation", "phrasebook",
    "sentence_frames", "grammar", "exercise",
}

# Trường bắt buộc cho từng loại section (ngoài "type")
REQUIRED_BY_TYPE = {
    "glossary": ["heading_vi", "heading_en", "items"],
    "parallel_text": ["heading_vi", "heading_en", "pairs"],
    "translation": ["heading_vi", "heading_en", "direction", "units"],
    "phrasebook": ["heading_vi", "heading_en", "items"],
    "sentence_frames": ["heading_vi", "heading_en", "items"],
    "grammar": ["point_en", "structure", "explain_vi"],
    "exercise": ["heading_vi", "heading_en", "questions"],
}


def validate(data):
    if not isinstance(data, dict):
        raise ValidationError("JSON gốc phải là một object (dấu ngoặc nhọn {...}).")

    # mode
    mode = data.get("mode")
    if mode not in ("clil", "english"):
        raise ValidationError(
            f"Trường 'mode' phải là 'clil' hoặc 'english', nhưng nhận được: {mode!r}."
        )

    for field in ("title_vi", "title_en"):
        if not data.get(field):
            raise ValidationError(f"Thiếu trường bắt buộc ở cấp tài liệu: '{field}'.")

    sections = data.get("sections")
    if not isinstance(sections, list) or len(sections) == 0:
        raise ValidationError("Trường 'sections' phải là một mảng có ít nhất 1 phần tử.")

    for idx, sec in enumerate(sections, start=1):
        if not isinstance(sec, dict):
            raise ValidationError(f"Section thứ {idx} không phải là object.")
        stype = sec.get("type")
        if stype not in VALID_TYPES:
            raise ValidationError(
                f"Section thứ {idx} có 'type' không hợp lệ: {stype!r}. "
                f"Chỉ chấp nhận: {', '.join(sorted(VALID_TYPES))}."
            )
        for field in REQUIRED_BY_TYPE[stype]:
            if field not in sec or sec[field] in (None, ""):
                raise ValidationError(
                    f"Section thứ {idx} (type={stype}) thiếu trường bắt buộc: '{field}'."
                )

        # Kiểm tra sâu các mảng con bắt buộc
        if stype == "glossary":
            _require_item_fields(sec["items"], ["vi", "en"], idx, stype, "items")
        elif stype == "parallel_text":
            _require_item_fields(sec["pairs"], ["vi", "en"], idx, stype, "pairs")
        elif stype == "translation":
            if sec.get("direction") not in ("vi2en", "en2vi"):
                raise ValidationError(
                    f"Section thứ {idx} (type=translation) có 'direction' không hợp lệ: "
                    f"{sec.get('direction')!r}. Chỉ chấp nhận 'vi2en' hoặc 'en2vi'."
                )
            _require_item_fields(sec["units"], ["source", "target"], idx, stype, "units")
        elif stype == "phrasebook":
            _require_item_fields(sec["items"], ["en", "vi"], idx, stype, "items")
        elif stype == "sentence_frames":
            _require_item_fields(sec["items"], ["frame_en", "frame_vi"], idx, stype, "items")
        elif stype == "exercise":
            _require_item_fields(sec["questions"], ["q_en"], idx, stype, "questions")

    return mode


def _require_item_fields(arr, fields, sec_idx, stype, arr_name):
    if not isinstance(arr, list) or len(arr) == 0:
        raise ValidationError(
            f"Section thứ {sec_idx} (type={stype}) có '{arr_name}' rỗng hoặc không phải mảng."
        )
    for j, item in enumerate(arr, start=1):
        if not isinstance(item, dict):
            raise ValidationError(
                f"Section thứ {sec_idx} (type={stype}), phần tử {j} của '{arr_name}' không phải object."
            )
        for f in fields:
            if f not in item or item[f] in (None, ""):
                raise ValidationError(
                    f"Section thứ {sec_idx} (type={stype}), phần tử {j} của '{arr_name}' "
                    f"thiếu trường bắt buộc: '{f}'."
                )


# ---------------------------------------------------------------------------
# HEADER TÀI LIỆU
# ---------------------------------------------------------------------------
def build_header(doc, data):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, data["title_vi"].upper(), size=15, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, data["title_en"], size=13, italic=True)

    # Dòng thông tin subject / grade / cefr
    info_parts = []
    if data.get("subject"):
        info_parts.append(str(data["subject"]))
    if data.get("grade"):
        info_parts.append(f"Lớp {data['grade']}")
    if data.get("cefr"):
        info_parts.append(f"CEFR {data['cefr']}")
    if info_parts:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p3, " • ".join(info_parts), size=12, italic=True)

    doc.add_paragraph()  # khoảng trống


def build_section_heading(doc, heading_vi, heading_en):
    p = doc.add_paragraph()
    add_run(p, heading_vi, size=13, bold=True)
    if heading_en:
        add_run(p, f"  ({heading_en})", size=12, bold=True, italic=True)
    return p


# ---------------------------------------------------------------------------
# SECTION: glossary
# ---------------------------------------------------------------------------
def _build_glossary_table(doc, items):
    """Dựng một bảng glossary cho danh sách items, bỏ cột rỗng toàn bộ."""
    col_defs = [
        ("vi", "Tiếng Việt"),
        ("en", "English"),
        ("ipa", "IPA"),
        ("pos", "Loại từ"),
        ("forms", "Dạng khác"),
        ("note_vi", "Ghi chú"),
    ]
    active_cols = []
    for key, label in col_defs:
        if key in ("vi", "en"):
            active_cols.append((key, label))  # luôn có 2 cột này
        else:
            if any(str(it.get(key, "")).strip() for it in items):
                active_cols.append((key, label))

    table = doc.add_table(rows=1, cols=len(active_cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0].cells
    for c, (key, label) in enumerate(active_cols):
        _fill_cell(hdr[c], [{"text": label, "bold": True, "size": 12}],
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(hdr[c], GRAY_FILL)

    # Dữ liệu
    for it in items:
        row = table.add_row().cells
        for c, (key, label) in enumerate(active_cols):
            val = str(it.get(key, "")).strip()
            italic = key == "ipa"
            _fill_cell(row[c], [{"text": val, "size": 12, "italic": italic}])


def build_glossary(doc, sec):
    build_section_heading(doc, sec["heading_vi"], sec["heading_en"])
    items = sec["items"]

    if sec.get("group_by") == "pos":
        # Nhóm theo từ loại, giữ thứ tự xuất hiện của các pos
        groups = {}
        order = []
        for it in items:
            key = str(it.get("pos", "")).strip() or "—"
            if key not in groups:
                groups[key] = []
                order.append(key)
            groups[key].append(it)
        for key in order:
            sub = doc.add_paragraph()
            sub.paragraph_format.space_before = Pt(4)
            label = POS_LABELS.get(key, key)
            add_run(sub, label, size=12, bold=True, italic=True)
            _build_glossary_table(doc, groups[key])
            doc.add_paragraph()
        return

    _build_glossary_table(doc, items)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# SECTION: parallel_text
# ---------------------------------------------------------------------------
def build_parallel_text(doc, sec):
    build_section_heading(doc, sec["heading_vi"], sec["heading_en"])
    pairs = sec["pairs"]
    numbered = bool(sec.get("numbered", False))
    granularity = sec.get("granularity", "sentence")
    # paragraph → giãn cách hàng rộng hơn
    row_space = Pt(8) if granularity == "paragraph" else Pt(2)

    ncols = 3 if numbered else 2
    table = doc.add_table(rows=0, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    for i, pair in enumerate(pairs, start=1):
        row = table.add_row().cells
        col = 0
        if numbered:
            p_num = _fill_cell(row[0], [{"text": f"{i}.", "bold": True, "size": 12}],
                               align=WD_ALIGN_PARAGRAPH.RIGHT)
            p_num.paragraph_format.space_after = row_space
            col = 1
        p_vi = _fill_cell(row[col], [{"text": str(pair.get("vi", ""))}])
        p_vi.paragraph_format.space_after = row_space
        p_en = _fill_cell(row[col + 1], [{"text": str(pair.get("en", "")), "italic": True}])
        p_en.paragraph_format.space_after = row_space

    # Cân cột
    for row in table.rows:
        if numbered:
            row.cells[0].width = Cm(1)
            row.cells[1].width = Cm(7.5)
            row.cells[2].width = Cm(7.5)
        else:
            row.cells[0].width = Cm(8)
            row.cells[1].width = Cm(8)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# SECTION: translation
# ---------------------------------------------------------------------------
def build_translation(doc, sec):
    build_section_heading(doc, sec["heading_vi"], sec["heading_en"])
    units = sec["units"]
    direction = sec["direction"]
    show_literal = bool(sec.get("show_literal", False))

    # Nhãn chiều dịch
    if direction == "vi2en":
        src_label, tgt_label = "VN", "EN"
    else:  # en2vi
        src_label, tgt_label = "EN", "VN"

    for i, unit in enumerate(units, start=1):
        if i > 1:
            add_thin_rule(doc)

        # Số thứ tự + source in đậm
        ps = doc.add_paragraph()
        add_run(ps, f"{i}. ", bold=True)
        add_run(ps, f"[{src_label}] ", bold=True, size=11, color=GRAY_TEXT)
        add_run(ps, str(unit["source"]), bold=True)

        # target thụt vào một cấp
        pt = doc.add_paragraph()
        pt.paragraph_format.left_indent = Cm(0.6)
        add_run(pt, f"[{tgt_label}] ", size=11, color=GRAY_TEXT)
        add_run(pt, str(unit["target"]))

        # literal
        if show_literal and str(unit.get("literal", "")).strip():
            pl = doc.add_paragraph()
            pl.paragraph_format.left_indent = Cm(0.6)
            add_run(pl, "Nghĩa đen: ", italic=True, size=11, color=GRAY_TEXT)
            add_run(pl, str(unit["literal"]), italic=True, size=11, color=GRAY_TEXT)

        # tense — chip nhỏ
        if str(unit.get("tense", "")).strip():
            ptn = doc.add_paragraph()
            ptn.paragraph_format.left_indent = Cm(0.6)
            add_run(ptn, f"Thì: {unit['tense']}", italic=True, size=11)

        # structures
        for st in unit.get("structures", []) or []:
            psr = doc.add_paragraph()
            psr.paragraph_format.left_indent = Cm(0.6)
            add_mono_run(psr, str(st.get("pattern", "")), bold=True, size=11)
            if st.get("name_vi"):
                add_run(psr, f"  ({st['name_vi']})", italic=True, size=11)
            if st.get("explain_vi"):
                pex = doc.add_paragraph()
                pex.paragraph_format.left_indent = Cm(1.0)
                add_run(pex, str(st["explain_vi"]), size=11)

        # vocab — bảng nhỏ gọn
        vocab = unit.get("vocab") or []
        if vocab:
            build_vocab_table(doc, vocab, size=11)

        # phrases — bullet list
        phrases = unit.get("phrases") or []
        if phrases:
            build_phrase_bullets(doc, phrases, indent=Cm(0.6), size=11)

        # note_vi
        if str(unit.get("note_vi", "")).strip():
            pn = doc.add_paragraph()
            pn.paragraph_format.left_indent = Cm(0.6)
            add_run(pn, f"📝 {unit['note_vi']}", italic=True, size=11)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# SECTION: phrasebook
# ---------------------------------------------------------------------------
PHRASEBOOK_COLS = [
    ("en", "English"),
    ("ipa", "IPA"),
    ("vi", "Tiếng Việt"),
    ("kind", "Loại"),
    ("example", "Ví dụ"),
]


def _build_phrasebook_table(doc, items):
    """Dựng một bảng phrasebook. Cột: English | IPA | Tiếng Việt | Loại | Ví dụ.

    Ví dụ (example_en + example_vi) gộp vào một ô, en đậm vi nghiêng.
    Bỏ cột IPA / Loại / Ví dụ nếu rỗng toàn bộ.
    """
    def has_example(it):
        return bool(str(it.get("example_en", "")).strip()
                    or str(it.get("example_vi", "")).strip())

    active = []
    for key, label in PHRASEBOOK_COLS:
        if key in ("en", "vi"):
            active.append((key, label))
        elif key == "example":
            if any(has_example(it) for it in items):
                active.append((key, label))
        elif any(str(it.get(key, "")).strip() for it in items):
            active.append((key, label))

    table = doc.add_table(rows=1, cols=len(active))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for c, (key, label) in enumerate(active):
        _fill_cell(hdr[c], [{"text": label, "bold": True, "size": 12}],
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(hdr[c], GRAY_FILL)

    for it in items:
        row = table.add_row().cells
        for c, (key, label) in enumerate(active):
            if key == "example":
                ex_en = str(it.get("example_en", "")).strip()
                ex_vi = str(it.get("example_vi", "")).strip()
                cell = row[c]
                cell.text = ""
                p0 = cell.paragraphs[0]
                if ex_en:
                    add_run(p0, ex_en, bold=True, size=11)
                if ex_vi:
                    pv = cell.add_paragraph() if ex_en else p0
                    add_run(pv, ex_vi, italic=True, size=11)
            else:
                val = str(it.get(key, "")).strip()
                italic = key == "ipa"
                _fill_cell(row[c], [{"text": val, "size": 11, "italic": italic}])


def build_phrasebook(doc, sec):
    build_section_heading(doc, sec["heading_vi"], sec["heading_en"])
    items = sec["items"]
    group = bool(sec.get("group_by_kind", True))

    if group:
        # Nhóm theo kind, giữ thứ tự theo PHRASE_KIND_LABELS rồi tới kind lạ
        groups = {}
        order = []
        for it in items:
            key = str(it.get("kind", "")).strip() or "—"
            if key not in groups:
                groups[key] = []
                order.append(key)
            groups[key].append(it)
        # Sắp xếp theo thứ tự nhãn định nghĩa sẵn, phần còn lại giữ nguyên
        known = [k for k in PHRASE_KIND_LABELS if k in groups]
        rest = [k for k in order if k not in PHRASE_KIND_LABELS]
        for key in known + rest:
            sub = doc.add_paragraph()
            sub.paragraph_format.space_before = Pt(4)
            label = PHRASE_KIND_LABELS.get(key, key if key != "—" else "Khác")
            add_run(sub, label, size=12, bold=True, italic=True)
            _build_phrasebook_table(doc, groups[key])
            doc.add_paragraph()
        return

    _build_phrasebook_table(doc, items)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# SECTION: sentence_frames
# ---------------------------------------------------------------------------
def _render_sentence_frame(doc, it):
    p = doc.add_paragraph()
    add_run(p, "• ", bold=True)
    add_run(p, str(it["frame_en"]), bold=True)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    add_run(p2, str(it["frame_vi"]), italic=True, size=12)

    if it.get("example_en"):
        pe = doc.add_paragraph()
        pe.paragraph_format.left_indent = Cm(1.0)
        add_run(pe, f"Ví dụ: {it['example_en']}", size=12)
    if it.get("example_vi"):
        pev = doc.add_paragraph()
        pev.paragraph_format.left_indent = Cm(1.0)
        add_run(pev, str(it["example_vi"]), italic=True, size=12)


def build_sentence_frames(doc, sec):
    build_section_heading(doc, sec["heading_vi"], sec["heading_en"])
    items = sec["items"]

    # Nếu có nhiều item cùng function_vi thì nhóm dưới tiêu đề phụ
    has_function = any(str(it.get("function_vi", "")).strip() for it in items)
    if has_function:
        groups = {}
        order = []
        for it in items:
            key = str(it.get("function_vi", "")).strip() or "—"
            if key not in groups:
                groups[key] = []
                order.append(key)
            groups[key].append(it)
        for key in order:
            sub = doc.add_paragraph()
            sub.paragraph_format.space_before = Pt(4)
            label = key if key != "—" else "Khác"
            add_run(sub, label, size=12, bold=True, italic=True)
            for it in groups[key]:
                _render_sentence_frame(doc, it)
        doc.add_paragraph()
        return

    for it in items:
        _render_sentence_frame(doc, it)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# SECTION: grammar
# ---------------------------------------------------------------------------
def build_grammar(doc, sec):
    # Tiêu đề point_en (kèm point_vi trong ngoặc)
    p = doc.add_paragraph()
    title = str(sec["point_en"])
    add_run(p, title, size=13, bold=True)
    if sec.get("point_vi"):
        add_run(p, f"  ({sec['point_vi']})", size=12, bold=True, italic=True)

    # Dòng thông tin tense / cefr
    meta = []
    if str(sec.get("tense", "")).strip():
        meta.append(f"Thì: {sec['tense']}")
    if str(sec.get("cefr", "")).strip():
        meta.append(f"CEFR {sec['cefr']}")
    if meta:
        pm = doc.add_paragraph()
        add_run(pm, " • ".join(meta), italic=True, size=11)

    # Khung structure(s) in đậm nền xám: khẳng định / phủ định / nghi vấn
    struct_lines = [("(+)", sec.get("structure"))]
    if str(sec.get("structure_negative", "")).strip():
        struct_lines.append(("(–)", sec["structure_negative"]))
    if str(sec.get("structure_question", "")).strip():
        struct_lines.append(("(?)", sec["structure_question"]))
    for j, (mark, formula) in enumerate(struct_lines):
        ps = doc.add_paragraph()
        _set_para_shading(ps, GRAY_FILL)
        ps.paragraph_format.left_indent = Cm(0.2)
        add_run(ps, f"{mark} ", bold=True)
        add_mono_run(ps, str(formula), bold=True)

    # signal_words
    signals = sec.get("signal_words") or []
    if signals:
        pw = doc.add_paragraph()
        add_run(pw, "Dấu hiệu nhận biết: ", bold=True, size=12)
        add_run(pw, ", ".join(str(s) for s in signals), italic=True, size=12)

    # explain_vi
    pe = doc.add_paragraph()
    add_run(pe, str(sec["explain_vi"]))

    # examples
    for ex in sec.get("examples", []) or []:
        pex = doc.add_paragraph()
        pex.paragraph_format.left_indent = Cm(0.5)
        add_run(pex, "– ", bold=True)
        add_run(pex, str(ex.get("en", "")))
        add_run(pex, f"  → {ex.get('vi', '')}", italic=True, size=12)

    # pitfalls_vi
    pitfalls = sec.get("pitfalls_vi") or []
    if pitfalls:
        pw = doc.add_paragraph()
        add_run(pw, "⚠ Lỗi thường gặp:", bold=True)
        for pit in pitfalls:
            pb = doc.add_paragraph(style="List Bullet")
            add_run(pb, str(pit), size=12)

    # contrast_with — khung so sánh hai điểm ngữ pháp
    contrast = sec.get("contrast_with")
    if isinstance(contrast, dict) and contrast.get("point_en"):
        pc = doc.add_paragraph()
        pc.paragraph_format.space_before = Pt(4)
        add_run(pc, "So sánh với: ", bold=True)
        add_run(pc, str(contrast["point_en"]), bold=True, italic=True)

        # Hai cột: điểm này vs điểm kia
        ctab = doc.add_table(rows=1, cols=2)
        ctab.style = "Table Grid"
        ctab.alignment = WD_TABLE_ALIGNMENT.CENTER
        pair = contrast.get("example_pair") or {}
        hdr = ctab.rows[0].cells
        _fill_cell(hdr[0], [{"text": str(sec["point_en"]), "bold": True, "size": 11}],
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _fill_cell(hdr[1], [{"text": str(contrast["point_en"]), "bold": True, "size": 11}],
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(hdr[0], GRAY_FILL)
        _set_cell_shading(hdr[1], GRAY_FILL)
        if pair.get("this_en") or pair.get("other_en"):
            row = ctab.add_row().cells
            _fill_cell(row[0], [{"text": str(pair.get("this_en", "")), "size": 11, "italic": True}])
            _fill_cell(row[1], [{"text": str(pair.get("other_en", "")), "size": 11, "italic": True}])

        # difference_vi + explain_vi phía dưới
        if str(contrast.get("explain_vi", "")).strip():
            pcx = doc.add_paragraph()
            add_run(pcx, str(contrast["explain_vi"]), size=11)
        if str(pair.get("difference_vi", "")).strip():
            pdf = doc.add_paragraph()
            add_run(pdf, f"Khác nhau: {pair['difference_vi']}", italic=True, size=11)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# SECTION: exercise
# ---------------------------------------------------------------------------
def build_exercise(doc, sec, counter, answers_collector):
    """counter: đối tượng đếm số câu xuyên suốt tài liệu (dict {'n': int}).
    answers_collector: list gom (số câu, answer, explain_vi) cho phần đáp án.
    """
    build_section_heading(doc, sec["heading_vi"], sec["heading_en"])

    if sec.get("instruction_en") or sec.get("instruction_vi"):
        pi = doc.add_paragraph()
        if sec.get("instruction_en"):
            add_run(pi, str(sec["instruction_en"]), italic=True, size=12)
        if sec.get("instruction_vi"):
            if sec.get("instruction_en"):
                add_run(pi, "  ", size=12)
            add_run(pi, f"({sec['instruction_vi']})", italic=True, size=12)

    letters = ["A", "B", "C", "D", "E", "F", "G", "H"]
    task = str(sec.get("task", "")).strip()

    # matching → bảng nối cột 2 cột A/B cho cả nhóm câu hỏi
    if task == "matching":
        _build_matching(doc, sec, counter, answers_collector, letters)
        doc.add_paragraph()
        return

    is_translate = task in ("translate_vi2en", "translate_en2vi")

    for q in sec["questions"]:
        counter["n"] += 1
        n = counter["n"]

        pq = doc.add_paragraph()
        add_run(pq, f"Câu {n}. ", bold=True)
        add_run(pq, str(q["q_en"]))

        if q.get("q_vi"):
            pv = doc.add_paragraph()
            pv.paragraph_format.left_indent = Cm(0.6)
            add_run(pv, str(q["q_vi"]), italic=True, size=11)

        choices = q.get("choices") or []
        if choices:
            for i, ch in enumerate(choices):
                pc = doc.add_paragraph()
                pc.paragraph_format.left_indent = Cm(0.6)
                letter = letters[i] if i < len(letters) else str(i + 1)
                add_run(pc, f"{letter}. ", bold=True)
                add_run(pc, str(ch))
        elif is_translate:
            # Câu dịch: chừa 2 dòng kẻ trống cho học sinh viết bản dịch
            for _ in range(2):
                pblank = doc.add_paragraph()
                pblank.paragraph_format.left_indent = Cm(0.6)
                add_run(pblank, "." * 90, size=12)
        else:
            # Câu điền/tự luận: chừa dòng kẻ trống
            pblank = doc.add_paragraph()
            pblank.paragraph_format.left_indent = Cm(0.6)
            add_run(pblank, "." * 90, size=12)

        # Gom đáp án
        answers_collector.append(
            {
                "n": n,
                "answer": q.get("answer", ""),
                "explain_vi": q.get("explain_vi", ""),
            }
        )

    doc.add_paragraph()


def _build_matching(doc, sec, counter, answers_collector, letters):
    """Bài nối cột: cột A là các q_en (đánh số), cột B là các choices (đánh chữ).

    Cột B gom toàn bộ choices theo thứ tự xuất hiện, loại trùng.
    """
    questions = sec["questions"]

    # Gom cột B
    col_b = []
    for q in questions:
        for ch in q.get("choices") or []:
            if str(ch) not in col_b:
                col_b.append(str(ch))

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    _fill_cell(hdr[0], [{"text": "Cột A", "bold": True, "size": 12}],
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_cell(hdr[1], [{"text": "Cột B", "bold": True, "size": 12}],
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_shading(hdr[0], GRAY_FILL)
    _set_cell_shading(hdr[1], GRAY_FILL)

    nrows = max(len(questions), len(col_b))
    for i in range(nrows):
        row = table.add_row().cells
        if i < len(questions):
            counter["n"] += 1
            n = counter["n"]
            q = questions[i]
            specs = [{"text": f"{n}. ", "bold": True, "size": 12},
                     {"text": str(q.get("q_en", "")), "size": 12}]
            _fill_cell(row[0], specs)
            answers_collector.append(
                {
                    "n": n,
                    "answer": q.get("answer", ""),
                    "explain_vi": q.get("explain_vi", ""),
                }
            )
        else:
            _fill_cell(row[0], [{"text": "", "size": 12}])
        if i < len(col_b):
            letter = letters[i] if i < len(letters) else str(i + 1)
            _fill_cell(row[1], [{"text": f"{letter}. ", "bold": True, "size": 12},
                                {"text": col_b[i], "size": 12}])
        else:
            _fill_cell(row[1], [{"text": "", "size": 12}])


# ---------------------------------------------------------------------------
# ĐÁP ÁN
# ---------------------------------------------------------------------------
def build_answer_key(doc, answers):
    doc.add_section(WD_SECTION.NEW_PAGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "ĐÁP ÁN — ANSWER KEY", size=14, bold=True)
    doc.add_paragraph()

    for a in answers:
        pa = doc.add_paragraph()
        add_run(pa, f"Câu {a['n']}. ", bold=True)
        ans = str(a.get("answer", "")).strip()
        add_run(pa, ans if ans else "(chưa có đáp án)")
        if a.get("explain_vi"):
            pe = doc.add_paragraph()
            pe.paragraph_format.left_indent = Cm(0.6)
            add_run(pe, str(a["explain_vi"]), italic=True, size=12)


# ---------------------------------------------------------------------------
# DỰNG TÀI LIỆU
# ---------------------------------------------------------------------------
def build(data):
    validate(data)

    doc = Document()
    setup_document(doc)
    build_header(doc, data)

    counter = {"n": 0}
    answers = []

    dispatch = {
        "glossary": lambda s: build_glossary(doc, s),
        "parallel_text": lambda s: build_parallel_text(doc, s),
        "translation": lambda s: build_translation(doc, s),
        "phrasebook": lambda s: build_phrasebook(doc, s),
        "sentence_frames": lambda s: build_sentence_frames(doc, s),
        "grammar": lambda s: build_grammar(doc, s),
        "exercise": lambda s: build_exercise(doc, s, counter, answers),
    }

    for sec in data["sections"]:
        dispatch[sec["type"]](sec)

    # answer_key default = True theo schema
    answer_key = data.get("answer_key", True)
    if answer_key and answers:
        build_answer_key(doc, answers)

    return doc


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------
def main(argv=None):
    parser = argparse.ArgumentParser(
        description="Dựng tài liệu học tập song ngữ Việt–Anh (.docx) từ JSON."
    )
    parser.add_argument("input", help="Đường dẫn file JSON đầu vào (đúng schema).")
    parser.add_argument("-o", "--output", required=True, help="Đường dẫn file .docx xuất ra.")
    args = parser.parse_args(argv)

    # Đọc JSON
    try:
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
    except FileNotFoundError:
        sys.stderr.write(f"LỖI: Không tìm thấy file đầu vào: {args.input}\n")
        return 1
    except json.JSONDecodeError as e:
        sys.stderr.write(
            f"LỖI: File JSON không hợp lệ ({args.input}).\n"
            f"Chi tiết: dòng {e.lineno}, cột {e.colno}: {e.msg}\n"
        )
        return 1

    # Validate + dựng
    try:
        doc = build(data)
    except ValidationError as e:
        sys.stderr.write(f"LỖI DỮ LIỆU: {e}\n")
        return 1

    try:
        doc.save(args.output)
    except OSError as e:
        sys.stderr.write(f"LỖI: Không ghi được file xuất ra: {e}\n")
        return 1

    print(f"Đã tạo: {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
