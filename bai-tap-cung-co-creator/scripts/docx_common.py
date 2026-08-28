# -*- coding: utf-8 -*-
"""Tiện ích dựng file .docx cho skill bai-tap-cung-co-creator.

Rút gọn từ exam-idea/scripts/docx_common.py, bỏ các hàm chỉ phục vụ bảng ma
trận / bảng đặc tả. Giữ nguyên quy ước font Times New Roman và cách ép rFonts
để tiếng Việt không bị đổi font khi mở trên máy khác.
"""

from __future__ import annotations

import json
import os
import sys


def force_utf8_stdout():
    """Bảo đảm in được tiếng Việt trên VPS Linux chạy locale C/POSIX."""
    for stream in (sys.stdout, sys.stderr):
        enc = (getattr(stream, "encoding", None) or "").lower()
        if enc.replace("-", "") not in ("utf8", "utf_8") and hasattr(stream, "reconfigure"):
            try:
                stream.reconfigure(encoding="utf-8", errors="replace")
            except (ValueError, OSError):
                pass


force_utf8_stdout()

from docx import Document  # noqa: E402
from docx.enum.section import WD_ORIENT  # noqa: E402
from docx.enum.table import WD_ALIGN_VERTICAL  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt, RGBColor  # noqa: E402

FONT = "Times New Roman"

PAGE_W_CM = 21.0          # A4 dọc
PAGE_H_CM = 29.7
MARGIN_CM = 2.0

ALIGN = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# --------------------------------------------------------------- low level
def _rfonts(run):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(attr), FONT)


def set_borders(table, sz="6"):
    pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    pr.append(borders)


def set_fixed_layout(table):
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def shade(cell, hexcolor="F2F2F2"):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    tcPr.append(el)


def write(cell, text="", bold=False, italic=False, align="center",
          size=10, valign="center"):
    """Ghi nội dung vào ô bảng. Xuống dòng bằng ký tự \\n."""
    cell.vertical_alignment = {
        "center": WD_ALIGN_VERTICAL.CENTER,
        "top": WD_ALIGN_VERTICAL.TOP,
        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
    }[valign]
    p = cell.paragraphs[0]
    p.alignment = ALIGN[align]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if text not in (None, ""):
        for i, line in enumerate(str(text).split("\n")):
            if i:
                p.add_run().add_break()
            r = p.add_run(line)
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(size)
            _rfonts(r)
    return cell


def apply_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def new_doc(landscape=False, margin=MARGIN_CM, base_size=12):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(base_size)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(0)
    sec = doc.sections[0]
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Cm(PAGE_H_CM), Cm(PAGE_W_CM)
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width, sec.page_height = Cm(PAGE_W_CM), Cm(PAGE_H_CM)
    sec.left_margin = sec.right_margin = Cm(margin)
    sec.top_margin = sec.bottom_margin = Cm(margin)
    return doc


def para(doc, text="", bold=False, italic=False, align="left", size=12,
         space_before=0, space_after=3, indent=0.0, color=None):
    p = doc.add_paragraph()
    p.alignment = ALIGN[align]
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.left_indent = Cm(indent)
    for i, line in enumerate(str(text).split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        _rfonts(r)
    return p


def rich_para(doc, chunks, align="left", size=12, space_before=0,
              space_after=3, indent=0.0):
    """chunks: list các tuple (text, {bold, italic, size, color})."""
    p = doc.add_paragraph()
    p.alignment = ALIGN[align]
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.left_indent = Cm(indent)
    for text, opt in chunks:
        r = p.add_run(str(text))
        r.bold = opt.get("bold", False)
        r.italic = opt.get("italic", False)
        r.font.size = Pt(opt.get("size", size))
        if opt.get("color"):
            r.font.color.rgb = RGBColor.from_string(opt["color"])
        _rfonts(r)
    return p


def title(doc, text, size=15, space_after=8):
    return para(doc, text, bold=True, align="center", size=size,
                space_after=space_after)


def heading(doc, text, size=13, space_before=10, space_after=4, color="1F4E79"):
    return para(doc, text, bold=True, align="left", size=size,
                space_before=space_before, space_after=space_after, color=color)


def page_break(doc):
    doc.add_page_break()


# ------------------------------------------------------------------ IO
def save_doc(doc, path):
    """Lưu docx, tự tạo thư mục cha nếu chưa có (Linux không tự tạo)."""
    path = os.path.abspath(path)
    parent = os.path.dirname(path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    doc.save(path)
    return path


def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def dump_json(obj, path):
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)


# ------------------------------------------------------- hằng số nghiệp vụ
NHOM = {
    "A": "Củng cố lý thuyết",
    "B": "Nhận diện – phân loại – so sánh",
    "C": "Tính toán theo công thức",
    "D": "Đọc và xử lí dữ liệu",
    "E": "Thí nghiệm – thực hành",
    "F": "Giải thích hiện tượng",
    "G": "Ứng dụng thực tiễn",
    "H": "Phát hiện và sửa lỗi sai",
    "I": "Tổng hợp – nâng cao",
}

MUC_DO = {
    "NB": "Nhận biết",
    "TH": "Thông hiểu",
    "VD": "Vận dụng",
    "VDC": "Vận dụng cao",
}

TANG = {
    1: ("★", "Cơ bản"),
    2: ("★★", "Chuẩn"),
    3: ("★★★", "Vận dụng"),
    4: ("★★★★", "Nâng cao"),
}


def sao(tang):
    """1 -> '★ (Cơ bản)'."""
    k, ten = TANG.get(int(tang), ("?", "?"))
    return f"{k} ({ten})"
