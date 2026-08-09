#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
normalize_format.py — Chuẩn hóa định dạng cho MỌI file giáo án .docx, chống lỗi:
  - Xóa đoạn trống/rác, số trang đứng một mình.
  - Tách các dòng bị dính (xuống dòng \n trong đoạn; hai gạch đầu dòng "  - " dồn một dòng).
  - Thụt lề ý chính (-) và ý phụ (+ / •) bằng INDENT TREO (hanging) — không dùng dấu cách/tab.
  - In đậm tiêu đề (TIẾT / I.II.III / HOẠT ĐỘNG / 1.2.3 / a)b) / Bước / "Sản phẩm cần đạt").
  - Đồng bộ font Times New Roman; hàng tiêu đề bảng in đậm, căn giữa, tô nền nhạt.

Dùng cho file tự sinh HOẶC file nhập từ ngoài (vd convert từ PDF/Word khác).

Cách dùng:
    python normalize_format.py input.docx [output.docx]
(không truyền output -> ghi đè input)
"""
import sys, re
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

# Console Windows mặc định là cp1252/cp437 -> in tiếng Việt sẽ ném
# UnicodeEncodeError. Ép UTF-8 để chạy giống nhau trên Windows lẫn Linux.
for _s in (sys.stdout, sys.stderr):
    try:
        _s.reconfigure(encoding="utf-8", errors="replace")
    except (AttributeError, ValueError):
        pass

FONT = "Times New Roman"
ROMAN = r'(?:I|II|III|IV|V|VI|VII|VIII|IX|X)'
HEADER_NUM = re.compile(r'^\d+\.\s*(Về|Kiến thức|Năng lực|Phẩm chất|Mục tiêu|Phương pháp|Công cụ|Giáo viên|Học sinh)')
UP_TITLE = re.compile(r'^(PHIẾU|MẪU|PHỤ LỤC|BẢNG|TIÊU CHÍ|RUBRIC|THIẾT KẾ BÀI DẠY|TÊN DỰ ÁN|ĐIỀU CHỈNH)')
UPPER = 'A-ZÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬĐÈÉẺẼẸÊẾỀỂỄỆÌÍỈĨỊÒÓỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÙÚỦŨỤƯỨỪỬỮỰỲÝỶỸỴ'
SPLIT_DASH = re.compile(r'\s-\s(?=[' + UPPER + r'])')


def set_font(r, size, bold=False, italic=False):
    r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    rpr = r._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(a), FONT)


def classify(s):
    if re.fullmatch(r'\d{1,3}', s):
        return None  # số trang -> bỏ
    if re.match(r'^TIẾT\b', s): return 'tiet'
    if re.match(r'^HOẠT ĐỘNG', s): return 'hd'
    if re.match(r'^' + ROMAN + r'\.\s', s): return 'roman'
    if UP_TITLE.match(s): return 'title'
    if re.match(r'^Bước\s*\d', s): return 'buoc'
    if re.match(r'^(Sản phẩm cần đạt|Nhiệm vụ học tập|Đáp án|Sản phẩm cuối dự án|Quy định chung|Lưu ý)', s): return 'spbold'
    if re.match(r'^[a-dA-D]\)\s', s): return 'ab'
    if HEADER_NUM.match(s): return 'subnum'
    if re.match(r'^Câu\s*\d', s): return 'cau'
    if s.startswith('- '): return 'dash'
    if s.startswith('+'): return 'plus'
    if s[:1] in '•▪◦‣·': return 'bullet'
    if re.match(r'^\d+[.)]\s?', s): return 'numlist'
    letters = [c for c in s if c.isalpha()]
    if letters and len(s) > 6 and s == s.upper(): return 'banner'
    return 'body'


def style_para(p, kind, in_table=False):
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(1 if in_table else 2)
    pf.line_spacing = 1.1 if in_table else 1.15
    pf.left_indent = Cm(0); pf.first_line_indent = Cm(0); pf.alignment = AL.LEFT
    size = 11.5 if in_table else 13; bold = False; italic = False
    if in_table:
        L1l, L1h, L2l, L2h = Cm(0.45), Cm(0.3), Cm(0.85), Cm(0.3)
    else:
        L1l, L1h, L2l, L2h = Cm(0.8), Cm(0.4), Cm(1.5), Cm(0.45)
    if kind == 'tiet':   bold = True; size += 1; pf.alignment = AL.CENTER; pf.space_before = Pt(14); pf.space_after = Pt(4)
    elif kind == 'banner': bold = True; pf.alignment = AL.CENTER; pf.space_after = Pt(6)
    elif kind == 'title':  bold = True; pf.alignment = AL.LEFT if in_table else AL.CENTER; pf.space_before = Pt(6); pf.space_after = Pt(3)
    elif kind == 'roman':  bold = True; pf.space_before = Pt(8); pf.space_after = Pt(3)
    elif kind == 'hd':     bold = True; pf.space_before = Pt(8); pf.space_after = Pt(3)
    elif kind == 'subnum': bold = True; pf.space_before = Pt(3)
    elif kind == 'ab':     bold = True; pf.left_indent = Cm(0.3)
    elif kind == 'buoc':   bold = True
    elif kind == 'spbold': bold = True
    elif kind == 'cau':    bold = True; pf.space_before = Pt(2)
    elif kind in ('dash', 'numlist'): pf.left_indent = L1l; pf.first_line_indent = -L1h
    elif kind in ('plus', 'bullet'):  pf.left_indent = L2l; pf.first_line_indent = -L2h
    for r in p.runs:
        set_font(r, size, bold, italic)


def to_lines(text):
    out = []
    for raw in text.split('\n'):
        s = raw.replace('\t', ' ').strip()
        s = re.sub(r'\s{2,}', ' ', s)
        if not s or classify(s) is None:
            continue
        # tách hai gạch đầu dòng bị dồn một dòng
        if s.startswith('- ') and SPLIT_DASH.search(s):
            for seg in SPLIT_DASH.split(s):
                seg = seg.strip()
                if seg:
                    out.append(seg if seg.startswith('- ') else '- ' + seg)
        else:
            out.append(s)
    return out


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def normalize(path_in, path_out=None):
    d = docx.Document(path_in)
    body = d.element.body
    for child in list(body):
        if child.tag != qn('w:p'):
            continue
        p = Paragraph(child, d)
        lines = to_lines(p.text)
        if not lines:
            body.remove(child); continue
        for s in lines:
            np = p.insert_paragraph_before('')
            np.add_run(s)
            style_para(np, classify(s), in_table=False)
        body.remove(child)
    # bìa: 2 đoạn đầu căn giữa, in đậm
    for p in d.paragraphs[:2]:
        p.paragraph_format.alignment = AL.CENTER
        for r in p.runs:
            r.font.bold = True
    # bảng
    for t in d.tables:
        for ri, row in enumerate(t.rows):
            for cell in row.cells:
                alllines = []
                for p in cell.paragraphs:
                    alllines += to_lines(p.text)
                for p in list(cell.paragraphs):
                    p._element.getparent().remove(p._element)
                if not alllines:
                    alllines = ['']
                for s in alllines:
                    np = cell.add_paragraph(); np.add_run(s)
                    if s == '':
                        for r in np.runs: set_font(r, 11.5)
                        np.paragraph_format.space_after = Pt(0)
                    else:
                        style_para(np, classify(s), in_table=True)
                    if ri == 0:
                        np.paragraph_format.alignment = AL.CENTER
                        for r in np.runs: r.font.bold = True
            if ri == 0:
                for cell in row.cells:
                    shade(cell, 'E8EEF7')
    d.save(path_out or path_in)
    print("Normalized ->", path_out or path_in)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python normalize_format.py input.docx [output.docx]")
        sys.exit(1)
    normalize(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
