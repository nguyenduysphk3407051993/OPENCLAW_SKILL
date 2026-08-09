#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_giaoan.py — Sinh GIÁO ÁN STEM (.docx) từ file JSON.

Hỗ trợ 2 cấu trúc (chọn qua trường "format"):
  - "5512"  (MẶC ĐỊNH): cấu trúc Kế hoạch bài dạy theo Công văn 5512, thiết kế chuyên nghiệp
            (trang bìa banner màu, thanh tiêu đề màu, hộp callout An toàn/Mẹo STEM/Nhiệm vụ/Tình huống,
            bảng lộ trình & phân vai). Mỗi tiết có I. MỤC TIÊU / II. THIẾT BỊ / III. TIẾN TRÌNH
            (bảng 2 cột GV–HS, 4 bước, "Sản phẩm cần đạt") / IV. ĐÁNH GIÁ / V. PHỤ LỤC.
  - "edp"   : cấu trúc cũ theo quy trình EDP (A. Mục đích → E. Cách thức tổ chức + bảng thử nghiệm).

Cách dùng:
    python build_giaoan.py noi_dung.json "Giao_an_STEM_<TenDuAn>.docx"

Yêu cầu: pip install python-docx --break-system-packages
"""
import sys, json
from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Console Windows mặc định là cp1252/cp437 -> in tiếng Việt sẽ ném
# UnicodeEncodeError. Ép UTF-8 để chạy giống nhau trên Windows lẫn Linux.
for _s in (sys.stdout, sys.stderr):
    try:
        _s.reconfigure(encoding="utf-8", errors="replace")
    except (AttributeError, ValueError):
        pass

FONT = "Times New Roman"

# Bảng màu (concept "xanh sinh học")
C_BANNER   = "1B5E20"   # xanh lá đậm — banner bìa
C_BAR      = "E8F5E9"   # nền thanh tiêu đề mục
C_BAR_TX   = "1B5E20"   # chữ thanh tiêu đề mục
C_ACT      = "F1F8E9"   # nền thanh hoạt động
C_ACT_TX   = "33691E"
C_HDR_FILL = "C8E6C9"   # nền hàng tiêu đề bảng
C_GROUP    = "F2F2F2"   # nền hàng tiêu đề nhóm trong rubric

# Callout: (nhãn, nền, màu chữ nhãn)
CALLOUTS = {
    "mission":   ("NHIỆM VỤ BIỆT ĐỘI", "E3F2FD", "0D47A1"),
    "safety":    ("AN TOÀN",           "FFEBEE", "B71C1C"),
    "tip":       ("MẸO STEM",          "FFF8E1", "E65100"),
    "situation": ("TÌNH HUỐNG",        "F3E5F5", "4A148C"),
    "note":      ("LƯU Ý",             "ECEFF1", "37474F"),
    "badge":     ("HUY HIỆU",          "FFF3E0", "BF360C"),
}

_USABLE_W = Emu(int(6.27 * 914400))   # ~ A4 trừ lề 2.5cm hai bên; cập nhật trong build()

# ----------------------------------------------------------------------------
def set_font(run, size=13, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ép đúng font cho mọi bảng mã (chống fallback Calibri làm lệch định dạng)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(a), FONT)

def _rgb(hexstr):
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))

# ---------------------------------------------------------------------------
# CHỐNG LỖI ĐỊNH DẠNG: làm sạch văn bản trước khi đưa vào paragraph/cell.
# - Bỏ tab, xuống dòng lạc, dấu cách thừa (KHÔNG dùng dấu cách/tab để thụt lề).
# - Bỏ marker thủ công người dùng có thể lỡ gõ ("- ", "+ ", "• ", "– ", "* ")
#   vì việc thụt lề & dấu đầu dòng do renderer tự đặt bằng indent treo.
import re as _re
_LEAD_MARK = _re.compile(r'^\s*[\-–•◦‣·\*\+]\s+')

def clean_text(s, strip_marker=True):
    if s is None:
        return ""
    s = str(s).replace('\t', ' ').replace('\r', ' ').replace('\n', ' ')
    s = _re.sub(r'\s{2,}', ' ', s).strip()
    if strip_marker:
        s = _LEAD_MARK.sub('', s)
    return s

def _hanging(p, level=1, in_table=False):
    """Thụt lề ý chính/ý phụ bằng indent treo (hanging) để dòng xuống dòng thẳng hàng."""
    if in_table:
        base = Pt(11) if level == 1 else Pt(22)
        hang = Pt(10)
    else:
        base = Pt(18) if level == 1 else Pt(34)
        hang = Pt(12)
    pf = p.paragraph_format
    pf.left_indent = base
    pf.first_line_indent = -hang

def finalize_doc(doc):
    """Bước rà cuối: ép đúng font mọi run (kể cả trong bảng), bỏ run rỗng,
    cắt khoảng trắng thừa đầu/cuối run — đảm bảo file xuất ra không lệch định dạng."""
    def _fix_paragraph(p):
        for r in p.runs:
            if r.text:
                r.text = r.text.replace('\t', ' ').replace('\r', '')
            rpr = r._element.get_or_add_rPr()
            rf = rpr.get_or_add_rFonts()
            if not rf.get(qn('w:ascii')):
                for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                    rf.set(qn(a), FONT)
    for p in doc.paragraphs:
        _fix_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _fix_paragraph(p)

def para(doc, text="", size=13, bold=False, italic=False, align=None,
         space_after=4, space_before=0, color=None, left_indent=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if left_indent:
        p.paragraph_format.left_indent = Pt(left_indent)
    # Giữ nguyên công thức/đáp án có dấu chấm câu, chỉ bỏ tab/xuống dòng & marker lạc
    r = p.add_run(clean_text(text, strip_marker=False))
    set_font(r, size, bold, italic, color)
    return p

def bullets(doc, items, size=13, mark="– ", indent=16):
    """Render danh sách gạch đầu dòng, thụt lề bằng indent treo (không dùng dấu cách).
    Mỗi phần tử có thể là:
      - chuỗi  -> ý chính (cấp 1);
      - dict {"text": "...", "sub": ["...", ...]} -> ý chính kèm ý phụ (cấp 2);
      - chuỗi mở đầu bằng "+"/"•" -> tự nhận là ý phụ (cấp 2).
    """
    for it in items or []:
        # Ý chính kèm ý phụ
        if isinstance(it, dict):
            txt = clean_text(it.get("text", ""))
            if txt:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                _hanging(p, level=1)
                set_font(p.add_run(mark + txt), size)
            for sub in it.get("sub", []):
                stxt = clean_text(sub)
                if not stxt:
                    continue
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(1)
                _hanging(sp, level=2)
                set_font(sp.add_run("+ " + stxt), size)
            continue
        raw = str(it)
        is_sub = bool(_re.match(r'^\s*[+•◦‣]', raw))
        txt = clean_text(raw)
        if not txt:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1 if is_sub else 2)
        _hanging(p, level=2 if is_sub else 1)
        set_font(p.add_run(("+ " if is_sub else mark) + txt), size)

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def _set_tbl_full_width(table):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(int(_USABLE_W.emu / 635)))   # 1 twip = 635 EMU

def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none')
        borders.append(e)
    tblPr.append(borders)

def _fill_cell(cell, items, bold=False, italic=False, align=None, size=13, color=None):
    if isinstance(items, str):
        items = [items]
    cell.text = ""
    first = True
    for it in items:
        txt = clean_text(it, strip_marker=False)
        if not txt and not first:
            continue
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(1)
        if align:
            p.alignment = align
        r = p.add_run(txt)
        set_font(r, size, bold=bold, italic=italic, color=color)

# ----- Khối thiết kế: banner, thanh tiêu đề, callout -------------------------
def banner(doc, lines):
    """Bìa: hộp full-width nền xanh đậm, chữ trắng căn giữa.
    lines: list (text, size, bold)."""
    t = doc.add_table(rows=1, cols=1)
    _set_tbl_full_width(t)
    _no_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    _shade(c, C_BANNER)
    c.text = ""
    first = True
    for (txt, size, bold) in lines:
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(txt)
        set_font(r, size, bold=bold, color=(0xFF, 0xFF, 0xFF))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def bar(doc, text, fill=C_BAR, txcolor=C_BAR_TX, size=13, space_before=6):
    """Thanh tiêu đề mục: hộp full-width nền nhạt, chữ đậm màu."""
    if space_before:
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(space_before)
    t = doc.add_table(rows=1, cols=1)
    _set_tbl_full_width(t)
    _no_borders(t)
    c = t.rows[0].cells[0]
    _shade(c, fill)
    txc = (int(txcolor[0:2], 16), int(txcolor[2:4], 16), int(txcolor[4:6], 16))
    _fill_cell(c, text, bold=True, size=size, color=txc)

def callout(doc, loai, noi_dung, tieu_de=None):
    label, fill, txc = CALLOUTS.get(loai, CALLOUTS["note"])
    if tieu_de:
        label = f"{label} — {tieu_de}"
    t = doc.add_table(rows=1, cols=1)
    _set_tbl_full_width(t)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    _shade(c, fill)
    c.text = ""
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label); set_font(r, 12, bold=True,
        color=(int(txc[0:2],16), int(txc[2:4],16), int(txc[4:6],16)))
    items = noi_dung if isinstance(noi_dung, list) else [noi_dung]
    for it in items:
        pp = c.add_paragraph(); pp.paragraph_format.space_after = Pt(1)
        rr = pp.add_run(it); set_font(rr, 12.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def make_table(doc, headers, rows, header_fill=C_HDR_FILL, center_cols=None, full=True):
    cols = len(headers)
    t = doc.add_table(rows=1, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if full:
        _set_tbl_full_width(t)
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _fill_cell(hdr[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(hdr[i], header_fill)
    if center_cols is None:
        center_cols = [0, cols - 1] if cols >= 2 else [0]
    for row in rows:
        cells = t.add_row().cells
        if len(row) == 1:
            merged = cells[0]
            for j in range(1, cols):
                merged = merged.merge(cells[j])
            _fill_cell(merged, row[0], bold=True, italic=True)
            _shade(merged, C_GROUP)
            continue
        is_total = str(row[0]).strip().lower().startswith(("tổng", "tong"))
        for i in range(cols):
            val = row[i] if i < len(row) else ""
            align = WD_ALIGN_PARAGRAPH.CENTER if i in center_cols else None
            _fill_cell(cells[i], val, bold=is_total, align=align)
            if is_total:
                _shade(cells[i], C_BAR)
    return t

# ----------------------------------------------------------------------------
def render_cover(doc, data):
    info = []
    if data.get("truong"):
        para(doc, data["truong"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, size=12)
    if data.get("to"):
        para(doc, data["to"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, size=12)
    if data.get("giao_vien"):
        para(doc, data["giao_vien"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, size=12)
    lines = [("KẾ HOẠCH BÀI DẠY STEM", 13, True),
             (data["ten_du_an"], 18, True)]
    if data.get("concept"):
        lines.append((data["concept"], 12.5, False))
    banner(doc, lines)

def render_mota(doc, data):
    bar(doc, "I. MÔ TẢ CHỦ ĐỀ", size=13)
    mt = data["mo_ta"]
    src = mt.get("thuc_trang")
    for d in (src if isinstance(src, list) else [src]):
        para(doc, d)
    if mt.get("dia_diem"):
        para(doc, "Địa điểm tổ chức: " + mt["dia_diem"])
    if mt.get("mon_chinh"):
        para(doc, "Môn học phụ trách chính: " + mt["mon_chinh"])
    for co in mt.get("callouts", []):
        callout(doc, co.get("loai", "note"), co["noi_dung"], co.get("tieu_de"))

def render_muc_tieu_chung(doc, mt):
    bar(doc, "II. MỤC TIÊU", size=13)
    if mt.get("kien_thuc"):
        para(doc, "1. Về kiến thức", bold=True, space_after=2)
        bullets(doc, mt["kien_thuc"])
    if mt.get("nang_luc_chung") or mt.get("nang_luc_khtn"):
        para(doc, "2. Về năng lực", bold=True, space_after=2)
        if mt.get("nang_luc_chung"):
            para(doc, "a) Năng lực chung", bold=True, italic=True, space_after=2, left_indent=10)
            bullets(doc, mt["nang_luc_chung"])
        if mt.get("nang_luc_khtn"):
            para(doc, "b) Năng lực khoa học tự nhiên", bold=True, italic=True, space_after=2, left_indent=10)
            bullets(doc, mt["nang_luc_khtn"])
    if mt.get("pham_chat"):
        para(doc, "3. Về phẩm chất", bold=True, space_after=2)
        bullets(doc, mt["pham_chat"])
    if mt.get("san_pham_stem"):
        para(doc, "4. Mục tiêu đối với sản phẩm STEM", bold=True, space_after=2)
        bullets(doc, mt["san_pham_stem"])

def render_meta_tables(doc, data):
    if data.get("phan_vai"):
        pv = data["phan_vai"]
        bar(doc, pv.get("tieu_de", "BẢNG PHÂN VAI BIỆT ĐỘI"), fill=C_ACT, txcolor=C_ACT_TX, size=12.5)
        make_table(doc, pv["headers"], pv["rows"])
    if data.get("timeline"):
        tl = data["timeline"]
        bar(doc, tl.get("tieu_de", "LỘ TRÌNH DỰ ÁN"), fill=C_ACT, txcolor=C_ACT_TX, size=12.5)
        make_table(doc, tl["headers"], tl["rows"])

def render_muc_tieu_tiet(doc, mt):
    bar(doc, "I. MỤC TIÊU", size=12.5)
    if mt.get("kien_thuc"):
        para(doc, "1. Kiến thức", bold=True, space_after=2)
        bullets(doc, mt["kien_thuc"])
    if mt.get("nang_luc_chung") or mt.get("nang_luc_khtn"):
        para(doc, "2. Năng lực", bold=True, space_after=2)
        if mt.get("nang_luc_chung"):
            para(doc, "a) Năng lực chung", bold=True, italic=True, space_after=2, left_indent=10)
            bullets(doc, mt["nang_luc_chung"])
        if mt.get("nang_luc_khtn"):
            para(doc, "b) Năng lực khoa học tự nhiên", bold=True, italic=True, space_after=2, left_indent=10)
            bullets(doc, mt["nang_luc_khtn"])
    if mt.get("pham_chat"):
        para(doc, "3. Phẩm chất", bold=True, space_after=2)
        bullets(doc, mt["pham_chat"])

def render_thiet_bi(doc, tb):
    bar(doc, "II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU", size=12.5)
    if tb.get("giao_vien"):
        para(doc, "1. Giáo viên", bold=True, space_after=2)
        bullets(doc, tb["giao_vien"])
    if tb.get("hoc_sinh"):
        para(doc, "2. Học sinh", bold=True, space_after=2)
        bullets(doc, tb["hoc_sinh"])

def render_to_chuc_table(doc, buoc_list):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_tbl_full_width(t)
    hdr = t.rows[0].cells
    _fill_cell(hdr[0], "HOẠT ĐỘNG CỦA GIÁO VIÊN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_cell(hdr[1], "HOẠT ĐỘNG CỦA HỌC SINH", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade(hdr[0], C_HDR_FILL); _shade(hdr[1], C_HDR_FILL)
    for b in buoc_list:
        cells = t.add_row().cells
        merged = cells[0].merge(cells[1])
        _fill_cell(merged, b.get("buoc", ""), bold=True)
        _shade(merged, C_GROUP)
        cells = t.add_row().cells
        _fill_cell(cells[0], b.get("gv", []) or [""])
        _fill_cell(cells[1], b.get("hs", []) or [""])
    return t

def render_hoat_dong(doc, hd):
    ten = hd.get("ten", "")
    if hd.get("thoi_gian"):
        ten = ten + f"  ({hd['thoi_gian']})"
    bar(doc, ten, fill=C_ACT, txcolor=C_ACT_TX, size=12.5, space_before=4)
    if hd.get("muc_tieu"):
        para(doc, "a) Mục tiêu", bold=True, space_after=2)
        bullets(doc, hd["muc_tieu"])
    for co in hd.get("callouts_truoc", []):
        callout(doc, co.get("loai", "note"), co["noi_dung"], co.get("tieu_de"))
    para(doc, "b) Tổ chức thực hiện", bold=True, space_after=2)
    if hd.get("to_chuc"):
        render_to_chuc_table(doc, hd["to_chuc"])
    for co in hd.get("callouts", []):
        callout(doc, co.get("loai", "note"), co["noi_dung"], co.get("tieu_de"))
    if hd.get("san_pham"):
        para(doc, "Sản phẩm cần đạt:", bold=True, italic=True, space_after=2, space_before=2)
        bullets(doc, hd["san_pham"])

def render_danh_gia_tiet(doc, dg):
    bar(doc, "IV. ĐÁNH GIÁ TRONG TIẾT HỌC", size=12.5)
    if dg.get("phuong_phap"):
        para(doc, "1. Phương pháp đánh giá", bold=True, space_after=2)
        bullets(doc, dg["phuong_phap"])
    if dg.get("cong_cu"):
        para(doc, "2. Công cụ đánh giá", bold=True, space_after=2)
        cc = dg["cong_cu"]
        if cc.get("mo_ta"):
            para(doc, cc["mo_ta"], space_after=2, italic=True)
        if cc.get("headers"):
            make_table(doc, cc["headers"], cc.get("rows", []))

def render_content_list(doc, items):
    """Render list noi_dung tổng quát: chuỗi / {bullets} / {text,bang} / {bang} / {callout}."""
    for muc in items:
        if isinstance(muc, dict):
            if muc.get("callout"):
                co = muc["callout"]
                callout(doc, co.get("loai", "note"), co["noi_dung"], co.get("tieu_de"))
                continue
            if muc.get("text"):
                para(doc, muc["text"], space_after=3)
            if muc.get("bullets"):
                bullets(doc, muc["bullets"])
            if muc.get("bang"):
                b = muc["bang"]
                if b.get("ten"):
                    para(doc, b["ten"], italic=True, space_after=2)
                make_table(doc, b["headers"], b["rows"])
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
        else:
            para(doc, muc, space_after=3)

def render_phieu_hoc_tap(doc, ph):
    bar(doc, ph["ten"], fill=C_ACT, txcolor=C_ACT_TX, size=12.5)
    if ph.get("thong_tin"):
        para(doc, ph["thong_tin"], italic=True, space_after=4)
    render_content_list(doc, ph.get("noi_dung", []))

def render_tiet(doc, tiet):
    doc.add_paragraph()
    banner(doc, [(tiet["ten_tiet"], 13, True),
                 (tiet.get("ten_bai", ""), 14, True)])
    if tiet.get("callouts_dau"):
        for co in tiet["callouts_dau"]:
            callout(doc, co.get("loai", "mission"), co["noi_dung"], co.get("tieu_de"))
    if tiet.get("muc_tieu"):
        render_muc_tieu_tiet(doc, tiet["muc_tieu"])
    if tiet.get("thiet_bi"):
        render_thiet_bi(doc, tiet["thiet_bi"])
    if tiet.get("hoat_dong"):
        bar(doc, "III. TIẾN TRÌNH DẠY HỌC", size=12.5)
        for i, hd in enumerate(tiet["hoat_dong"], 1):
            if "ten" not in hd or not hd["ten"].strip().upper().startswith("HOẠT ĐỘNG"):
                hd = dict(hd, ten=f"HOẠT ĐỘNG {i}. " + hd.get("ten", ""))
            render_hoat_dong(doc, hd)
    if tiet.get("danh_gia"):
        render_danh_gia_tiet(doc, tiet["danh_gia"])
    if tiet.get("phu_luc_tiet"):
        bar(doc, "V. PHỤ LỤC", size=12.5)
        for ph in tiet["phu_luc_tiet"]:
            render_phieu_hoc_tap(doc, ph)

def render_danh_gia_chi_tiet(doc, dgct):
    doc.add_page_break()
    banner(doc, [(dgct.get("tieu_de", "TIÊU CHÍ ĐÁNH GIÁ CHI TIẾT"), 15, True)])
    for rb in dgct.get("rubric", []):
        bar(doc, rb["ten"], fill=C_ACT, txcolor=C_ACT_TX, size=12.5)
        make_table(doc, rb.get("headers", ["STT", "Tiêu chí", "Điểm tối đa"]), rb["rows"])
    if dgct.get("tong_hop"):
        th = dgct["tong_hop"]
        bar(doc, th.get("ten", "PHIẾU TỔNG HỢP KẾT QUẢ"), size=12.5)
        make_table(doc, th.get("headers", ["Nội dung đánh giá", "Điểm tối đa"]), th["rows"])

def render_kien_thuc_lien_quan(doc, ktlq, tieu_de="PHỤ LỤC — CÁC KIẾN THỨC LIÊN QUAN"):
    doc.add_page_break()
    banner(doc, [(tieu_de, 15, True)])
    for kt in ktlq:
        bar(doc, kt["tieu_de"], size=12.5)
        render_content_list(doc, [d if isinstance(d, (str, dict)) else str(d) for d in kt["noi_dung"]])

def build_5512(data, out_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(13)
    global _USABLE_W
    sec = doc.sections[0]
    _USABLE_W = Emu(sec.page_width - sec.left_margin - sec.right_margin)

    render_cover(doc, data)
    render_mota(doc, data)
    if data.get("muc_tieu_chung"):
        render_muc_tieu_chung(doc, data["muc_tieu_chung"])
    render_meta_tables(doc, data)

    for tiet in data["tiet"]:
        render_tiet(doc, tiet)

    if data.get("danh_gia_chi_tiet"):
        render_danh_gia_chi_tiet(doc, data["danh_gia_chi_tiet"])
    if data.get("kien_thuc_lien_quan"):
        render_kien_thuc_lien_quan(doc, data["kien_thuc_lien_quan"])

    para(doc, "— Hết —", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8)
    finalize_doc(doc)
    doc.save(out_path)
    print("Saved (5512):", out_path)

# ----------------------------------------------------------------------------
# FORMAT EDP (giữ tương thích bản cũ)
# ----------------------------------------------------------------------------
def _edp_section(doc, text):
    para(doc, text, bold=True, space_after=2)

def thu_nghiem_block(doc, tieu_chi):
    para(doc, "Hướng dẫn chế tạo và thử nghiệm sản phẩm", bold=True, space_after=2)
    rows = [[str(i + 1), tc, ""] for i, tc in enumerate(tieu_chi)]
    para(doc, "Thử nghiệm lần 1", bold=True, italic=True, space_after=2)
    bullets(doc, ["Quan sát, ghi nhận đầy đủ các tiến trình và kết quả.",
                  "Đánh giá mức độ hoạt động của sản phẩm so với tiêu chí đã đặt ra ban đầu:"], mark="")
    make_table(doc, ["STT", "TIÊU CHÍ", "ĐẠT/KHÔNG ĐẠT"], rows)
    bullets(doc, ["Phần nào trong thiết kế hoạt động tốt?",
                  "Phần nào trong thiết kế hoạt động không tốt?",
                  "Có thể làm gì để cải tiến thiết kế của mình? Phác hoạ và ghi rõ."], mark="")
    para(doc, "Các lần thử nghiệm lần sau", bold=True, italic=True, space_after=2)
    bullets(doc, ["Các cải tiến đã thực hiện là gì? (lưu lại ảnh sản phẩm cải tiến)",
                  "Đánh giá mức độ hoạt động của sản phẩm so với tiêu chí đã đặt ra ban đầu:"], mark="")
    make_table(doc, ["STT", "TIÊU CHÍ", "ĐẠT/KHÔNG ĐẠT"],
               [[str(i + 1), tc, ""] for i, tc in enumerate(tieu_chi)])
    bullets(doc, ["Phần nào trong thiết kế hoạt động tốt?",
                  "Phần nào trong thiết kế hoạt động không tốt?",
                  "Có thể làm gì để cải tiến thiết kế của mình?",
                  "Điều chỉnh sản phẩm đến phiên bản tốt nhất trong điều kiện cho phép."], mark="")
    # (kết thúc khối thử nghiệm EDP)

def build_edp(data, out_path):
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = FONT; style.font.size = Pt(13)
    global _USABLE_W
    sec = doc.sections[0]
    _USABLE_W = Emu(sec.page_width - sec.left_margin - sec.right_margin)

    if data.get("truong"):
        para(doc, data["truong"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    if data.get("to"):
        para(doc, data["to"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    if data.get("giao_vien"):
        para(doc, data["giao_vien"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    para(doc, "THIẾT KẾ BÀI DẠY STEM", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x1F,0x49,0x7D))
    para(doc, "TÊN DỰ ÁN: " + data["ten_du_an"], size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x1F,0x49,0x7D))
    doc.add_paragraph()

    para(doc, "I. Mô tả chủ đề", bold=True, size=14)
    mt = data["mo_ta"]
    src = mt.get("thuc_trang")
    for d in (src if isinstance(src, list) else [src]):
        para(doc, d)
    if mt.get("dia_diem"): para(doc, "Địa điểm tổ chức: " + mt["dia_diem"])
    if mt.get("mon_chinh"): para(doc, "Môn học phụ trách chính: " + mt["mon_chinh"])

    para(doc, "II. Mục tiêu", bold=True, size=14)
    mu = data["muc_tieu"]
    _edp_section(doc, "1. Kỹ năng: Sau chủ đề, HS có khả năng"); bullets(doc, mu["ky_nang"])
    _edp_section(doc, "2. Thái độ"); bullets(doc, mu["thai_do"])
    _edp_section(doc, "3. Về định hướng phát triển năng lực"); bullets(doc, mu["nang_luc"])
    _edp_section(doc, "4. Thiết bị"); bullets(doc, mu["thiet_bi"])
    _edp_section(doc, "5. Tiến trình dạy học")

    for tiet in data["tiet"]:
        doc.add_paragraph()
        para(doc, tiet["ten_tiet"], size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        para(doc, tiet["ten_hoat_dong"], size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x1F,0x49,0x7D))
        _edp_section(doc, "A. Mục đích")
        para(doc, "Sau hoạt động này, HS có khả năng:", italic=True, space_after=2)
        bullets(doc, tiet["muc_dich"])
        _edp_section(doc, "B. Nội dung")
        for d in tiet["noi_dung"]: para(doc, d)
        _edp_section(doc, "C. Dự kiến sản phẩm hoạt động của học sinh")
        bullets(doc, tiet["san_pham"])
        if tiet.get("chuan_bi"):
            _edp_section(doc, "D. CHUẨN BỊ")
            cb = tiet["chuan_bi"]
            para(doc, "1. Giáo viên", bold=True, space_after=2); bullets(doc, cb["giao_vien"])
            para(doc, "2. Học sinh", bold=True, space_after=2); bullets(doc, cb["hoc_sinh"])
        _edp_section(doc, "E. Cách thức tổ chức hoạt động")
        for d in tiet.get("to_chuc", []): para(doc, d)
        if tiet.get("che_tao_thu_nghiem"):
            ctn = tiet["che_tao_thu_nghiem"]
            para(doc, "Chế tạo: " + ctn.get("che_tao", ""), space_after=2)
            thu_nghiem_block(doc, ctn["tieu_chi"])

    doc.add_page_break()
    para(doc, "PHỤ LỤC", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x1F,0x49,0x7D))
    pl = data.get("phu_luc", {})
    if pl.get("rubric"):
        para(doc, "Phụ lục 1. Các bảng tiêu chí đánh giá", bold=True, size=14)
        for rb in pl["rubric"]:
            para(doc, rb["ten"], bold=True, italic=True, space_after=2)
            make_table(doc, ["TT", "Tiêu chí", "Điểm"], rb["rows"])
    if pl.get("kien_thuc_lien_quan"):
        para(doc, "Phụ lục 2. Các kiến thức liên quan", bold=True, size=14)
        for kt in pl["kien_thuc_lien_quan"]:
            para(doc, kt["tieu_de"], bold=True, space_after=2)
            for d in kt["noi_dung"]: para(doc, d)
    if pl.get("phieu_hoc_tap"):
        doc.add_page_break()
        para(doc, "PHIẾU HỌC TẬP", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
        for ph in pl["phieu_hoc_tap"]:
            para(doc, ph["ten"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
            if ph.get("thong_tin"): para(doc, ph["thong_tin"], italic=True, space_after=4)
            render_content_list(doc, ph.get("noi_dung", []))

    para(doc, "— Hết —", align=WD_ALIGN_PARAGRAPH.CENTER)
    finalize_doc(doc)
    doc.save(out_path)
    print("Saved (edp):", out_path)

# ----------------------------------------------------------------------------
def build(data, out_path):
    fmt = data.get("format", "5512").lower()
    if fmt == "edp":
        build_edp(data, out_path)
    else:
        build_5512(data, out_path)

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python build_giaoan.py <input.json> <output.docx>")
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        data = json.load(f)
    build(data, sys.argv[2])
